"""
Eduxellence Export Engine — Phase 1
=====================================
Generates APA 7th Edition Word (.docx) and branded PDF reports
from any statistical analysis result dict.

Word:  python-docx  — italicised p-values, formatted tables, narrative text
PDF:   reportlab    — branded, publication-ready, chart-embedded

by Eduxellence Analytics · https://eduxellence.org
"""

import io, base64, re
from datetime import datetime

# ── Word / docx ────────────────────────────────────────────────────────────
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── PDF / reportlab ────────────────────────────────────────────────────────
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch, cm
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT, TA_JUSTIFY
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
    HRFlowable, Image, KeepTogether, PageBreak
)
from reportlab.platypus import BaseDocTemplate, PageTemplate, Frame
from reportlab.lib.pagesizes import A4

# ── Brand ──────────────────────────────────────────────────────────────────
BRAND_NAVY  = RGBColor(0x0B, 0x18, 0x29)
BRAND_BLUE  = RGBColor(0x1E, 0x6B, 0xFF)
BRAND_TEAL  = RGBColor(0x0F, 0xC9, 0xA0)
BRAND_SLATE = RGBColor(0x64, 0x74, 0x8B)

RL_NAVY  = colors.HexColor('#0B1829')
RL_BLUE  = colors.HexColor('#1E6BFF')
RL_TEAL  = colors.HexColor('#0FC9A0')
RL_SLATE = colors.HexColor('#64748B')
RL_LIGHT = colors.HexColor('#F7F9FC')
RL_BORDER= colors.HexColor('#E2E8F0')

SITE = "https://eduxellence.org"
NOW  = datetime.now().strftime("%d %B %Y, %H:%M")


# ══════════════════════════════════════════════════════════════════════════
#  HELPERS
# ══════════════════════════════════════════════════════════════════════════

def _strip_md(text: str) -> str:
    """Remove markdown bold markers."""
    return re.sub(r'\*\*(.+?)\*\*', r'\1', text or "")

def _add_hor_rule(doc: Document):
    """Add a thin horizontal rule paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'E2E8F0')
    pBdr.append(bottom)
    pPr.append(pBdr)

def _set_cell_bg(cell, hex_color: str):
    """Set cell background colour."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def _p_val_run(para, p_val):
    """Add an italicised, formatted p-value run to a paragraph."""
    if p_val is None or p_val == "—":
        para.add_run("—")
        return
    try:
        pf = float(str(p_val).replace("p","").replace("=","").replace("<","").strip())
        text = "< .001" if pf < 0.001 else f"= {pf:.3f}"
        run = para.add_run(f"p {text}")
        run.italic = True
        run.font.name = "Times New Roman"
    except Exception:
        run = para.add_run(str(p_val))
        run.italic = True


# ══════════════════════════════════════════════════════════════════════════
#  APA WORD DOCUMENT
# ══════════════════════════════════════════════════════════════════════════

def generate_docx(results: dict, filename_stem: str = "analysis") -> bytes:
    """
    Generate a fully APA 7th Edition formatted Word document.
    Returns bytes ready to stream as a download.
    """
    doc = Document()

    # ── Page margins (APA: 1 inch all sides) ──────────────────────────────
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1)
        section.right_margin  = Inches(1)

    # ── Default font ──────────────────────────────────────────────────────
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing    = Pt(24)  # double spacing
    style.paragraph_format.space_after     = Pt(0)
    style.paragraph_format.space_before    = Pt(0)

    # ── Cover block ───────────────────────────────────────────────────────
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_r = title_p.add_run("Statistical Analysis Report")
    title_r.bold      = True
    title_r.font.size = Pt(16)
    title_r.font.name = 'Times New Roman'
    title_r.font.color.rgb = BRAND_NAVY

    doc.add_paragraph()  # blank line

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_r = sub_p.add_run(
        f"{results.get('analysis_label', results.get('test', 'Analysis'))}"
    )
    sub_r.bold      = True
    sub_r.font.size = Pt(13)
    sub_r.font.color.rgb = BRAND_BLUE

    meta_p = doc.add_paragraph()
    meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_r = meta_p.add_run(
        f"Generated: {NOW}\nPowered by Eduxellence Analytics — {SITE}"
    )
    meta_r.font.size  = Pt(10)
    meta_r.font.color.rgb = BRAND_SLATE
    meta_r.font.italic = True

    _add_hor_rule(doc)
    doc.add_paragraph()

    # ── 1. Test Overview ───────────────────────────────────────────────────
    _apa_heading(doc, "1. Test Overview", level=1)
    overview_table = doc.add_table(rows=0, cols=2)
    overview_table.style = 'Table Grid'
    overview_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    overview_rows = [
        ("Statistical Test",   results.get("test", "—")),
        ("Sample Size (N)",    str(results.get("n", "—"))),
        ("p-value",            results.get("p_display", "—")),
        ("Significance",       results.get("significance", "—")),
        ("Effect Size",        _get_effect(results)),
        ("Generated",          NOW),
    ]
    for label, value in overview_rows:
        row = overview_table.add_row()
        _set_cell_bg(row.cells[0], '0B1829')
        lbl_run = row.cells[0].paragraphs[0].add_run(label)
        lbl_run.bold = True
        lbl_run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        lbl_run.font.size = Pt(11)
        val_run = row.cells[1].paragraphs[0].add_run(value)
        val_run.font.size = Pt(11)

    doc.add_paragraph()

    # ── 2. APA Citation ────────────────────────────────────────────────────
    if results.get("apa_citation"):
        _apa_heading(doc, "2. APA-Formatted Result", level=1)
        cite_p = doc.add_paragraph()
        cite_p.paragraph_format.left_indent  = Inches(0.5)
        cite_p.paragraph_format.right_indent = Inches(0.5)
        cite_r = cite_p.add_run(results["apa_citation"])
        cite_r.italic    = True
        cite_r.font.size = Pt(12)
        # Italicise p inside the citation
        doc.add_paragraph()

    # ── 3. Results Tables ─────────────────────────────────────────────────
    table_keys = [
        ("numeric_summary",   "3. Descriptive Statistics Table"),
        ("summary_table",     "3. Group Summary Table"),
        ("coef_table",        "3. Regression Coefficients"),
        ("pairs_table",       "3. Correlation Matrix"),
        ("posthoc_table",     "3. Post-Hoc Comparisons (Bonferroni)"),
    ]
    for key, heading in table_keys:
        if results.get(key):
            _apa_heading(doc, heading, level=1)
            _apa_table(doc, results[key])
            doc.add_paragraph()

    # Categorical frequency tables
    if results.get("categorical_summary"):
        for cs in results["categorical_summary"]:
            _apa_heading(doc, f"Frequency Table: {cs['variable']}", level=2)
            _apa_table(doc, cs["table"])
            doc.add_paragraph()

    # Contingency table
    if results.get("contingency_table"):
        _apa_heading(doc, "Contingency Table (Observed Frequencies)", level=2)
        _contingency_table_docx(doc, results["contingency_table"])
        doc.add_paragraph()

    # ── 4. Plain-English Interpretation ───────────────────────────────────
    if results.get("interpretation"):
        _apa_heading(doc, "4. Interpretation", level=1)
        clean = _strip_md(results["interpretation"])
        for para_text in clean.split("\n\n"):
            if para_text.strip():
                p = doc.add_paragraph()
                p.paragraph_format.first_line_indent = Inches(0.5)
                r = p.add_run(para_text.strip().replace("\n", " "))
                r.font.size = Pt(12)
        doc.add_paragraph()

    # ── 5. Methodology Note ───────────────────────────────────────────────
    _apa_heading(doc, "5. Methodology Note", level=1)
    method_p = doc.add_paragraph()
    method_p.paragraph_format.first_line_indent = Inches(0.5)
    method_text = (
        f"This analysis was conducted using the Eduxellence Analytics platform "
        f"({SITE}). Statistical computations were performed using Python's SciPy "
        f"library (scipy.stats). Charts were generated using Matplotlib and Seaborn. "
        f"Results are reported in accordance with APA 7th Edition guidelines. "
        f"Effect sizes and confidence intervals are included where applicable. "
        f"All analyses were run on {NOW}."
    )
    r = method_p.add_run(method_text)
    r.font.size = Pt(12)
    doc.add_paragraph()

    # ── Footer watermark ───────────────────────────────────────────────────
    _add_hor_rule(doc)
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_r = footer_p.add_run(
        f"Generated by Eduxellence Analytics · {SITE} · {NOW}"
    )
    footer_r.font.size   = Pt(9)
    footer_r.font.italic = True
    footer_r.font.color.rgb = BRAND_SLATE

    # ── Save & return ──────────────────────────────────────────────────────
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def _apa_heading(doc: Document, text: str, level: int = 1):
    """Add an APA-style bold centred (level 1) or bold flush-left (level 2) heading."""
    p = doc.add_paragraph()
    if level == 1:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.first_line_indent = Inches(0)
    r = p.add_run(text)
    r.bold      = True
    r.font.size = Pt(13) if level == 1 else Pt(12)
    r.font.name = 'Times New Roman'
    r.font.color.rgb = BRAND_NAVY
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(6)


def _apa_table(doc: Document, rows: list):
    """Render a list-of-dicts as an APA-formatted table."""
    if not rows:
        return
    cols = list(rows[0].keys())
    table = doc.add_table(rows=1, cols=len(cols))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    hdr = table.rows[0]
    for i, col in enumerate(cols):
        cell = hdr.cells[i]
        _set_cell_bg(cell, '0B1829')
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(str(col))
        run.bold            = True
        run.font.color.rgb  = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size       = Pt(10)
        run.font.name       = 'Times New Roman'

    # Data rows
    for i, row_data in enumerate(rows):
        row = table.add_row()
        if i % 2 == 1:
            for cell in row.cells:
                _set_cell_bg(cell, 'F7F9FC')
        for j, col in enumerate(cols):
            cell = row.cells[j]
            p    = cell.paragraphs[0]
            val  = str(row_data.get(col, ""))
            # Italicise p-value cells
            if col.lower() in ("p-value", "p (bonferroni)", "p_display", "p") and val not in ("—", ""):
                try:
                    pf = float(val)
                    run = p.add_run("< .001" if pf < 0.001 else f"= {pf:.3f}")
                    run.italic    = True
                    run.font.size = Pt(10)
                except Exception:
                    p.add_run(val).font.size = Pt(10)
            else:
                run = p.add_run(val)
                run.font.size = Pt(10)
                run.font.name = 'Times New Roman'
                # Bold significant rows
                if col in ("Significance", "Significant") and val in ("***", "**", "*", "Yes"):
                    run.bold = True
                    run.font.color.rgb = RGBColor(0x16, 0xA3, 0x4A)


def _contingency_table_docx(doc: Document, ct: dict):
    """Render a contingency table dict from pandas .to_dict()."""
    col_keys = list(ct.keys())
    if not col_keys:
        return
    row_keys = list(ct[col_keys[0]].keys())
    table = doc.add_table(rows=len(row_keys) + 1, cols=len(col_keys) + 1)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header
    _set_cell_bg(table.cell(0, 0), '0B1829')
    for j, ck in enumerate(col_keys):
        cell = table.cell(0, j + 1)
        _set_cell_bg(cell, '0B1829')
        run = cell.paragraphs[0].add_run(str(ck))
        run.bold           = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size      = Pt(10)

    # Rows
    for i, rk in enumerate(row_keys):
        rh = table.cell(i + 1, 0)
        _set_cell_bg(rh, 'F1F5F9')
        rh_run = rh.paragraphs[0].add_run(str(rk))
        rh_run.bold      = True
        rh_run.font.size = Pt(10)
        for j, ck in enumerate(col_keys):
            val = ct[ck].get(rk, "")
            table.cell(i + 1, j + 1).paragraphs[0].add_run(str(val)).font.size = Pt(10)


def _get_effect(results: dict) -> str:
    """Extract the most relevant effect size string."""
    for key in ("cohens_d", "cramers_v", "eta_squared", "r_squared", "effect_r"):
        if results.get(key) is not None:
            label = {"cohens_d": "Cohen's d", "cramers_v": "Cramér's V",
                     "eta_squared": "η²", "r_squared": "R²", "effect_r": "r"}[key]
            return f"{label} = {results[key]}"
    if results.get("effect_size"):
        return results["effect_size"]
    return "—"


# ══════════════════════════════════════════════════════════════════════════
#  PDF REPORT (reportlab)
# ══════════════════════════════════════════════════════════════════════════

def generate_pdf(results: dict, filename_stem: str = "analysis") -> bytes:
    """
    Generate a branded, publication-ready PDF report.
    Returns bytes ready to stream as a download.
    """
    buf = io.BytesIO()
    PAGE_W, PAGE_H = A4

    # ── Styles ─────────────────────────────────────────────────────────────
    styles = getSampleStyleSheet()

    title_style = ParagraphStyle('EduxTitle',
        fontName='Helvetica-Bold', fontSize=22, textColor=RL_NAVY,
        alignment=TA_CENTER, spaceAfter=6)

    sub_style = ParagraphStyle('EduxSub',
        fontName='Helvetica-Bold', fontSize=14, textColor=RL_BLUE,
        alignment=TA_CENTER, spaceAfter=4)

    meta_style = ParagraphStyle('EduxMeta',
        fontName='Helvetica-Oblique', fontSize=9, textColor=RL_SLATE,
        alignment=TA_CENTER, spaceAfter=16)

    h1_style = ParagraphStyle('EduxH1',
        fontName='Helvetica-Bold', fontSize=13, textColor=RL_NAVY,
        spaceBefore=14, spaceAfter=6, borderPad=4,
        borderColor=RL_BLUE, borderWidth=0, leftIndent=0)

    body_style = ParagraphStyle('EduxBody',
        fontName='Helvetica', fontSize=10, textColor=colors.HexColor('#1E293B'),
        leading=16, alignment=TA_JUSTIFY, spaceAfter=8, firstLineIndent=18)

    apa_style = ParagraphStyle('EduxAPA',
        fontName='Helvetica-Oblique', fontSize=10.5, textColor=RL_NAVY,
        leading=16, alignment=TA_JUSTIFY, spaceAfter=8,
        leftIndent=18, rightIndent=18,
        borderColor=colors.HexColor('#C4B5FD'), borderWidth=1,
        borderPad=8, backColor=colors.HexColor('#F8F7FF'))

    footer_style = ParagraphStyle('EduxFooter',
        fontName='Helvetica-Oblique', fontSize=8, textColor=RL_SLATE,
        alignment=TA_CENTER)

    # ── Build story ────────────────────────────────────────────────────────
    story = []

    # Cover header block
    story.append(Spacer(1, 0.3 * inch))
    story.append(Paragraph("Statistical Analysis Report", title_style))
    story.append(Paragraph(
        results.get("analysis_label", results.get("test", "Analysis")),
        sub_style))
    story.append(Paragraph(
        f"Generated: {NOW} &nbsp;·&nbsp; Powered by Eduxellence Analytics &nbsp;·&nbsp; {SITE}",
        meta_style))
    story.append(HRFlowable(width="100%", thickness=2, color=RL_BLUE,
                             spaceAfter=12))

    # 1. Test Overview Table
    story.append(Paragraph("1. Test Overview", h1_style))
    overview_data = [
        [_bold_pdf("Statistical Test"),   results.get("test", "—")],
        [_bold_pdf("Sample Size (N)"),    str(results.get("n", "—"))],
        [_bold_pdf("p-value"),            _fmt_p_pdf(results.get("p_value"))],
        [_bold_pdf("Significance"),       results.get("significance", "—")],
        [_bold_pdf("Effect Size"),        _get_effect(results)],
        [_bold_pdf("Generated"),          NOW],
    ]
    ov_table = Table(overview_data, colWidths=[2.2*inch, 4.3*inch])
    ov_table.setStyle(TableStyle([
        ('BACKGROUND',  (0,0), (0,-1), RL_NAVY),
        ('TEXTCOLOR',   (0,0), (0,-1), colors.white),
        ('FONTNAME',    (0,0), (0,-1), 'Helvetica-Bold'),
        ('FONTSIZE',    (0,0), (-1,-1), 10),
        ('BACKGROUND',  (1,0), (1,-1), RL_LIGHT),
        ('ROWBACKGROUNDS', (1,0), (1,-1), [RL_LIGHT, colors.white]),
        ('GRID',        (0,0), (-1,-1), 0.5, RL_BORDER),
        ('VALIGN',      (0,0), (-1,-1), 'MIDDLE'),
        ('LEFTPADDING', (0,0), (-1,-1), 8),
        ('RIGHTPADDING',(0,0), (-1,-1), 8),
        ('TOPPADDING',  (0,0), (-1,-1), 5),
        ('BOTTOMPADDING',(0,0), (-1,-1), 5),
    ]))
    story.append(ov_table)
    story.append(Spacer(1, 12))

    # 2. APA Citation
    if results.get("apa_citation"):
        story.append(Paragraph("2. APA-Formatted Result", h1_style))
        story.append(Paragraph(results["apa_citation"], apa_style))
        story.append(Spacer(1, 8))

    # 3. Results Tables
    table_keys = [
        ("numeric_summary",   "3. Descriptive Statistics"),
        ("summary_table",     "3. Group Summary"),
        ("coef_table",        "3. Regression Coefficients"),
        ("pairs_table",       "3. Correlation Pairs"),
        ("posthoc_table",     "3. Post-Hoc Comparisons"),
    ]
    for key, heading in table_keys:
        if results.get(key):
            story.append(Paragraph(heading, h1_style))
            story.append(_pdf_table(results[key]))
            story.append(Spacer(1, 10))

    if results.get("categorical_summary"):
        for cs in results["categorical_summary"]:
            story.append(Paragraph(f"Frequency Table: {cs['variable']}", h1_style))
            story.append(_pdf_table(cs["table"]))
            story.append(Spacer(1, 10))

    # 4. Interpretation
    if results.get("interpretation"):
        story.append(Paragraph("4. Interpretation", h1_style))
        clean = _strip_md(results["interpretation"])
        for para_text in clean.split("\n\n"):
            if para_text.strip():
                story.append(Paragraph(para_text.strip().replace("\n"," "), body_style))
        story.append(Spacer(1, 8))

    # 5. Charts (embed as images)
    if results.get("charts"):
        story.append(Paragraph("5. Charts & Visualisations", h1_style))
        for chart in results["charts"]:
            try:
                img_bytes = base64.b64decode(chart["img"])
                img_buf   = io.BytesIO(img_bytes)
                img = Image(img_buf, width=6.2*inch, height=3.8*inch)
                caption_style = ParagraphStyle('Caption',
                    fontName='Helvetica-Oblique', fontSize=9,
                    textColor=RL_SLATE, alignment=TA_CENTER, spaceAfter=10)
                story.append(KeepTogether([
                    img,
                    Paragraph(f"Figure. {chart.get('title','Chart')}", caption_style),
                    Spacer(1, 8)
                ]))
            except Exception:
                pass

    # 6. Methodology Note
    story.append(Paragraph("6. Methodology Note", h1_style))
    method_text = (
        f"This analysis was conducted using the Eduxellence Analytics platform ({SITE}). "
        f"Statistical computations used Python's SciPy library. Charts were generated "
        f"using Matplotlib and Seaborn. Results are reported following APA 7th Edition "
        f"guidelines. Effect sizes and 95% confidence intervals are included where applicable. "
        f"Analysis timestamp: {NOW}."
    )
    story.append(Paragraph(method_text, body_style))
    story.append(Spacer(1, 16))

    # Footer
    story.append(HRFlowable(width="100%", thickness=1, color=RL_BORDER, spaceAfter=6))
    story.append(Paragraph(
        f"Generated by Eduxellence Analytics · {SITE} · {NOW}",
        footer_style))

    # ── Build PDF ──────────────────────────────────────────────────────────
    doc = SimpleDocTemplate(
        buf, pagesize=A4,
        leftMargin=inch, rightMargin=inch,
        topMargin=inch,  bottomMargin=inch,
        title="Statistical Analysis Report — Eduxellence Analytics",
        author="Eduxellence Analytics",
        subject=results.get("test", "Statistical Analysis"),
    )
    doc.build(story)
    buf.seek(0)
    return buf.read()


def _bold_pdf(text: str):
    """Paragraph with bold text for PDF tables."""
    return Paragraph(f"<b>{text}</b>",
        ParagraphStyle('bold', fontName='Helvetica-Bold', fontSize=10,
                       textColor=colors.white))


def _fmt_p_pdf(p_val) -> str:
    """Format p-value for PDF display."""
    if p_val is None:
        return "—"
    try:
        pf = float(p_val)
        return "< .001" if pf < 0.001 else f"= {pf:.3f}"
    except Exception:
        return str(p_val)


def _pdf_table(rows: list) -> Table:
    """Convert list-of-dicts to a styled reportlab Table."""
    if not rows:
        return Spacer(1, 4)
    cols = list(rows[0].keys())
    max_cols = min(len(cols), 8)   # cap for wide tables
    vis_cols  = cols[:max_cols]

    # Header
    header = [Paragraph(f"<b><font color='white'>{c}</font></b>",
              ParagraphStyle('th', fontName='Helvetica-Bold', fontSize=9,
                             textColor=colors.white))
              for c in vis_cols]
    data = [header]

    # Rows
    for row_data in rows:
        row = []
        for col in vis_cols:
            val = str(row_data.get(col, ""))
            # Italicise p-values
            if col.lower() in ("p-value","p (bonferroni)","p_display","p"):
                try:
                    pf = float(val)
                    val = f"< .001" if pf < 0.001 else f"= {pf:.3f}"
                    row.append(Paragraph(f"<i>{val}</i>",
                        ParagraphStyle('italic', fontName='Helvetica-Oblique',
                                       fontSize=9, textColor=RL_NAVY)))
                    continue
                except Exception:
                    pass
            # Colour significant values
            if col in ("Significance","Significant") and val in ("***","**","*","Yes"):
                row.append(Paragraph(f"<b><font color='#16A34A'>{val}</font></b>",
                    ParagraphStyle('sig', fontName='Helvetica-Bold', fontSize=9)))
            else:
                row.append(Paragraph(val,
                    ParagraphStyle('td', fontName='Helvetica', fontSize=9,
                                   textColor=colors.HexColor('#1E293B'))))
        data.append(row)

    # Calculate equal col widths
    available_w = 6.5 * inch
    col_w = available_w / max_cols

    t = Table(data, colWidths=[col_w] * max_cols, repeatRows=1)
    t.setStyle(TableStyle([
        ('BACKGROUND',   (0,0), (-1,0),  RL_NAVY),
        ('ROWBACKGROUNDS',(0,1),(-1,-1), [RL_LIGHT, colors.white]),
        ('GRID',         (0,0), (-1,-1), 0.4, RL_BORDER),
        ('VALIGN',       (0,0), (-1,-1), 'MIDDLE'),
        ('LEFTPADDING',  (0,0), (-1,-1), 6),
        ('RIGHTPADDING', (0,0), (-1,-1), 6),
        ('TOPPADDING',   (0,0), (-1,-1), 4),
        ('BOTTOMPADDING',(0,0), (-1,-1), 4),
    ]))
    return t


# ══════════════════════════════════════════════════════════════════════════
#  ENHANCED EXCEL (already existed — upgraded with formatting)
# ══════════════════════════════════════════════════════════════════════════

def generate_excel(results: dict) -> bytes:
    """
    Generate a formatted Excel workbook with branded header row colours.
    """
    import pandas as pd
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter

    buf = io.BytesIO()
    with pd.ExcelWriter(buf, engine="openpyxl") as writer:
        # Sheet 1: Summary
        meta = {
            "Analysis":     results.get("test", ""),
            "Generated":    NOW,
            "Website":      SITE,
            "p-value":      results.get("p_display", "—"),
            "Significance": results.get("significance", "—"),
            "Effect Size":  _get_effect(results),
            "APA Citation": results.get("apa_citation", "—"),
            "Interpretation": results.get("interpretation", "—"),
        }
        pd.DataFrame([meta]).T.reset_index().rename(
            columns={"index": "Metric", 0: "Value"}
        ).to_excel(writer, sheet_name="Summary", index=False)

        # Sheet 2+: Results tables
        for key, sheet in [
            ("numeric_summary",  "Descriptive Stats"),
            ("summary_table",    "Group Summary"),
            ("coef_table",       "Coefficients"),
            ("pairs_table",      "Correlations"),
            ("posthoc_table",    "Post Hoc"),
        ]:
            if results.get(key):
                pd.DataFrame(results[key]).to_excel(
                    writer, sheet_name=sheet, index=False)

        if results.get("categorical_summary"):
            for cs in results["categorical_summary"]:
                pd.DataFrame(cs["table"]).to_excel(
                    writer, sheet_name=f"Freq_{cs['variable'][:18]}", index=False)

        # Apply header formatting to all sheets
        wb = writer.book
        navy_fill = PatternFill("solid", fgColor="0B1829")
        white_font = Font(bold=True, color="FFFFFF", name="Calibri", size=11)
        center_align = Alignment(horizontal="center", vertical="center", wrap_text=True)

        for ws in wb.worksheets:
            # Style header row
            for cell in ws[1]:
                cell.fill      = navy_fill
                cell.font      = white_font
                cell.alignment = center_align
            # Auto-fit columns
            for col in ws.columns:
                max_len = max((len(str(cell.value or "")) for cell in col), default=8)
                ws.column_dimensions[get_column_letter(col[0].column)].width = min(max_len + 4, 40)
            # Freeze top row
            ws.freeze_panes = "A2"

    buf.seek(0)
    return buf.read()
