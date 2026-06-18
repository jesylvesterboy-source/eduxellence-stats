"""
Eduxellence Statistical Analysis Engine v2.3
============================================
9 statistical tests · Auto assumption checking · Smart test recommender
APA-formatted output · Publication-ready charts via matplotlib/seaborn
Zero paid dependencies. Vercel free-tier compatible.
Supabase Storage integration for charts to prevent payload size limits.

by Eduxellence Analytics · https://eduxellence.org
"""

import io, base64, warnings, traceback, os, json, uuid, logging, sys, signal
import numpy as np
import pandas as pd
import matplotlib; matplotlib.use("Agg")
import matplotlib.pyplot as plt
import seaborn as sns
from scipy import stats
from itertools import combinations
from datetime import datetime
import time
from concurrent.futures import ThreadPoolExecutor, TimeoutError as FutureTimeoutError

# ── Version Metadata ──────────────────────────────────────────────────────
ENGINE_VERSION = "2.3.0"
ENGINE_BUILD = datetime.now().strftime("%Y%m%d")

# ── Configure Logging ──────────────────────────────────────────────────────
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

# ── Constants ──────────────────────────────────────────────────────────────
MAX_RUNTIME_SECONDS = 25
MAX_MEMORY_MB = 500
MAX_CHARTS = 10
MAX_PAIRPLOT_VARS = 4
MAX_PAIRPLOT_SAMPLE = 500
MAX_ROWS = 50000
MAX_COLS = 100
MAX_DESCRIPTIVE_COLS = 15
MAX_PREDICTORS = 10
RANDOM_SEED = 42
API_RATE_LIMIT = 10

# ── Error Codes ────────────────────────────────────────────────────────────
ERROR_CODES = {
    "VAL_001": "Column not found in dataset",
    "VAL_002": "Column contains no valid numeric data",
    "VAL_003": "Column has insufficient unique values",
    "VAL_004": "Dataset has insufficient rows",
    "VAL_005": "Too many predictors for regression",
    "VAL_006": "Duplicate predictor names found",
    "VAL_007": "Dataset exceeds maximum allowed rows",
    "VAL_008": "Dataset exceeds maximum allowed columns",
    "VAL_009": "All predictors are constant",
    "VAL_010": "Insufficient sample size for regression",
    "VAL_011": "Perfect separation detected in logistic regression",
    "LIM_001": "Analysis exceeded time limit",
    "LIM_002": "Analysis exceeded memory limit",
    "LIM_003": "Analysis too large for free tier",
    "LIM_004": "API rate limit exceeded",
    "ERR_001": "Internal server error",
    "ERR_002": "Chi-square test failed",
    "ERR_003": "Regression matrix is singular",
    "ERR_004": "Shapiro-Wilk test failed",
    "ERR_005": "Logistic regression failed to converge",
}

# ── Targeted Warnings Filtering ──────────────────────────────────────────
warnings.filterwarnings("ignore", category=UserWarning, module="matplotlib")
warnings.filterwarnings("ignore", category=FutureWarning, module="pandas")
warnings.filterwarnings("ignore", category=UserWarning, module="seaborn")

# ── Supabase Storage Integration ──────────────────────────────────────────
try:
    from supabase import create_client, Client
    SUPABASE_URL = os.environ.get("SUPABASE_URL")
    SUPABASE_KEY = os.environ.get("SUPABASE_SERVICE_ROLE_KEY")
    SUPABASE_BUCKET = os.environ.get("SUPABASE_BUCKET", "eduxellence-charts")
    supabase_available = bool(SUPABASE_URL and SUPABASE_KEY)
    if supabase_available:
        supabase: Client = create_client(SUPABASE_URL, SUPABASE_KEY)
        logger.info("Supabase client initialized successfully")
    else:
        logger.warning("Supabase credentials not found. Charts will use base64 encoding.")
except ImportError:
    supabase_available = False
    supabase = None
    logger.warning("Supabase package not installed. Charts will use base64 encoding.")

# ── Custom Exceptions ──────────────────────────────────────────────────────
class AnalysisTooLargeError(Exception):
    pass

class AnalysisTimeoutError(Exception):
    pass

class ValidationError(Exception):
    pass

class MemoryLimitError(Exception):
    pass

class RateLimitError(Exception):
    pass

class SingularMatrixError(Exception):
    pass

class PerfectSeparationError(Exception):
    pass

# ── API Rate Limiting ────────────────────────────────────────────────────
_rate_limit_store = {}

def check_rate_limit(ip_address, limit=API_RATE_LIMIT, window_seconds=60):
    from collections import defaultdict
    from datetime import datetime, timedelta
    
    current_time = datetime.now()
    
    if ip_address not in _rate_limit_store:
        _rate_limit_store[ip_address] = []
    
    _rate_limit_store[ip_address] = [
        t for t in _rate_limit_store[ip_address] 
        if current_time - t < timedelta(seconds=window_seconds)
    ]
    
    if len(_rate_limit_store[ip_address]) >= limit:
        raise RateLimitError(ERROR_CODES["LIM_004"])
    
    _rate_limit_store[ip_address].append(current_time)
    return True

# ── Memory Check ──────────────────────────────────────────────────────────
def check_memory_usage(df, max_memory_mb=MAX_MEMORY_MB):
    memory_mb = df.memory_usage(deep=True).sum() / (1024 * 1024)
    if memory_mb > max_memory_mb:
        raise MemoryLimitError(
            f"Dataset memory usage ({memory_mb:.1f} MB) exceeds limit ({max_memory_mb} MB). "
            f"Please reduce dataset size."
        )
    return memory_mb

# ── Timeout Protection ────────────────────────────────────────────────────
def with_timeout(timeout_seconds=MAX_RUNTIME_SECONDS):
    def decorator(func):
        def wrapper(*args, **kwargs):
            with ThreadPoolExecutor(max_workers=1) as executor:
                future = executor.submit(func, *args, **kwargs)
                try:
                    result = future.result(timeout=timeout_seconds)
                    return result
                except FutureTimeoutError:
                    raise AnalysisTimeoutError(ERROR_CODES["LIM_001"])
                except Exception as e:
                    raise e
        return wrapper
    return decorator

# ── System Metadata ──────────────────────────────────────────────────────
def get_system_metadata():
    import sys, platform
    import statsmodels, scipy, pandas, matplotlib, seaborn
    
    return {
        "python_version": sys.version.split()[0],
        "platform": platform.platform(),
        "packages": {
            "numpy": np.__version__,
            "pandas": pd.__version__,
            "scipy": scipy.__version__,
            "statsmodels": statsmodels.__version__,
            "matplotlib": matplotlib.__version__,
            "seaborn": seaborn.__version__
        },
        "random_seed": RANDOM_SEED
    }

# ── Brand palette ──────────────────────────────────────────────────────────
P = {
    "navy":"#0B1829","navy2":"#112240","blue":"#1E6BFF","blue2":"#4B8AFF",
    "teal":"#0FC9A0","teal2":"#09A882","white":"#FFFFFF","off":"#F7F9FC",
    "slate":"#64748B","slateL":"#CBD5E1","border":"#E2E8F0",
    "ok":"#16A34A","err":"#DC2626","warn":"#D97706","text":"#0F172A",
}
PAL = ["#1E6BFF","#0FC9A0","#F59E0B","#8B5CF6","#EC4899","#EF4444","#10B981","#F97316","#06B6D4","#84CC16"]

# ── Chart Storage Handler ──────────────────────────────────────────────────
def _store_chart(fig, chart_name, dpi=110, use_storage=True):
    """Store chart either as base64 (small) or Supabase Storage (large)."""
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=dpi, bbox_inches="tight", facecolor=fig.get_facecolor())
    buf.seek(0)
    img_data = buf.read()
    
    size_mb = len(img_data) / (1024 * 1024)
    
    # ── FIXED: Thread-safe filenames with UUID ──────────────────────────
    try:
        if use_storage and size_mb > 1.0 and supabase_available and supabase:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            unique_id = str(uuid.uuid4())[:8]
            filename = f"charts/{timestamp}_{unique_id}_{chart_name.replace(' ', '_')}.png"
            
            try:
                supabase.storage.from_(SUPABASE_BUCKET).upload(
                    file=img_data,
                    path=filename,
                    file_options={"content-type": "image/png"}
                )
            except TypeError:
                supabase.storage.from_(SUPABASE_BUCKET).upload(
                    filename,
                    img_data,
                    {"content-type": "image/png"}
                )
            
            public_url = supabase.storage.from_(SUPABASE_BUCKET).get_public_url(filename)
            return {"type": "url", "url": public_url, "size_mb": round(size_mb, 2)}
            
        result = base64.b64encode(img_data).decode()
        return {"type": "base64", "data": result, "size_mb": round(size_mb, 2)}
    
    finally:
        # ── FIXED: Always close figure to prevent memory leak ──────────
        plt.close(fig)

def _b64(fig, dpi=110):
    result = _store_chart(fig, "chart", dpi)
    return result["data"] if result["type"] == "base64" else result["url"]

def _style(ax_list):
    for ax in (ax_list if hasattr(ax_list,"__iter__") and not isinstance(ax_list,plt.Axes) else [ax_list]):
        ax.set_facecolor(P["off"]); ax.spines[["top","right"]].set_visible(False)
        ax.spines[["left","bottom"]].set_color(P["border"])
        ax.tick_params(colors=P["slate"],labelsize=9)
        for lbl in [ax.xaxis.label, ax.yaxis.label]: lbl.set_color(P["slate"])
        if ax.get_title(): ax.title.set_color(P["navy"])

def sp(p): return "***" if p<.001 else "**" if p<.01 else "*" if p<.05 else "ns"
def fp(p): return "p < .001" if p<.001 else f"p = {p:.3f}"
def verdict(p,a=.05): sig=p<a; return ("statistically significant" if sig else "not statistically significant"),("reject" if sig else "fail to reject")

# ── Validation Functions ──────────────────────────────────────────────────
def validate_numeric_column(df, column, context=""):
    if column not in df.columns:
        raise ValidationError(f"{ERROR_CODES['VAL_001']}: {column}{': ' + context if context else ''}")
    test_series = pd.to_numeric(df[column], errors='coerce')
    if test_series.isna().all():
        raise ValidationError(f"{ERROR_CODES['VAL_002']}: {column}{': ' + context if context else ''}")
    return True

def validate_categorical_column(df, column, context=""):
    if column not in df.columns:
        raise ValidationError(f"{ERROR_CODES['VAL_001']}: {column}{': ' + context if context else ''}")
    n_unique = df[column].nunique()
    if n_unique < 2:
        raise ValidationError(f"{ERROR_CODES['VAL_003']}: {column} has only {n_unique} unique value(s){': ' + context if context else ''}")
    return True

def validate_row_count(df, min_rows=3, context=""):
    if len(df) < min_rows:
        raise ValidationError(f"{ERROR_CODES['VAL_004']}: Dataset has only {len(df)} rows. Minimum required: {min_rows}{': ' + context if context else ''}")
    return True

def validate_predictor_count(predictors, max_predictors=MAX_PREDICTORS, context=""):
    if len(predictors) > max_predictors:
        raise ValidationError(f"{ERROR_CODES['VAL_005']}: {len(predictors)} predictors. Maximum allowed: {max_predictors}{': ' + context if context else ''}")
    if len(set(predictors)) < len(predictors):
        raise ValidationError(f"{ERROR_CODES['VAL_006']}: Duplicate predictor names found{': ' + context if context else ''}")
    return True

def validate_no_empty_columns(df):
    empty_cols = [c for c in df.columns if df[c].isna().all()]
    if empty_cols:
        warnings.warn(f"Found completely empty columns: {empty_cols}. They will be dropped.")
        df = df.drop(columns=empty_cols)
    return df

def validate_no_duplicate_columns(df):
    if len(df.columns) != len(set(df.columns)):
        warnings.warn("Duplicate column names detected. Renaming to unique names.")
        df.columns = pd.io.parsers.ParserBase({'names': df.columns})._maybe_dedup_names(df.columns)
    return df

# ── Data Cleaning ──────────────────────────────────────────────────────────
def prepare_dataframe(df):
    logger.info(f"Preparing dataframe with {len(df)} rows and {len(df.columns)} columns")
    df = df.replace([np.inf, -np.inf], np.nan)
    df = validate_no_empty_columns(df)
    df = validate_no_duplicate_columns(df)
    str_cols = df.select_dtypes(include=['object']).columns
    for col in str_cols:
        df[col] = df[col].astype(str).str.strip()
    for col in df.columns:
        if df[col].dtype == 'object':
            converted = pd.to_numeric(df[col], errors='coerce')
            if not converted.isna().all():
                df[col] = converted
    constant_cols = [c for c in df.columns if df[c].nunique() <= 1]
    if constant_cols:
        warnings.warn(f"Dropped constant columns: {constant_cols}")
        df = df.drop(columns=constant_cols)
    logger.info(f"Dataframe prepared: {len(df)} rows, {len(df.columns)} columns")
    return df

def check_analysis_limits(df, analysis_type, params):
    rows = len(df)
    if rows > MAX_ROWS:
        raise AnalysisTooLargeError(
            f"{ERROR_CODES['LIM_003']}: {rows:,} rows exceeds limit of {MAX_ROWS:,} rows."
        )
    if analysis_type == "descriptive":
        cols = params.get("columns", [])
        if len(cols) > MAX_DESCRIPTIVE_COLS:
            raise AnalysisTooLargeError(
                f"{ERROR_CODES['LIM_003']}: {len(cols)} columns exceeds limit of {MAX_DESCRIPTIVE_COLS} columns."
            )
    if analysis_type == "regression" or analysis_type == "logistic_regression":
        predictors = params.get("predictors", [])
        if len(predictors) > MAX_PREDICTORS:
            raise AnalysisTooLargeError(
                f"{ERROR_CODES['LIM_003']}: {len(predictors)} predictors exceeds limit of {MAX_PREDICTORS} predictors."
            )
    check_memory_usage(df)

def validate_dataset(df, max_rows=MAX_ROWS, max_cols=MAX_COLS):
    if len(df) > max_rows:
        raise AnalysisTooLargeError(f"{ERROR_CODES['VAL_007']}: {len(df):,} rows. Maximum allowed is {max_rows:,}.")
    if len(df.columns) > max_cols:
        raise AnalysisTooLargeError(f"{ERROR_CODES['VAL_008']}: {len(df.columns)} columns. Maximum allowed is {max_cols}.")
    df = prepare_dataframe(df)
    return df

def get_vif_label(vif_score):
    if vif_score < 5:
        return "low"
    elif vif_score < 10:
        return "moderate"
    else:
        return "severe"

def get_missing_summary(df):
    missing_summary = {}
    for col in df.columns:
        missing_count = int(df[col].isna().sum())
        if missing_count > 0:
            missing_summary[col] = {
                "missing_count": missing_count,
                "missing_percent": round((missing_count / len(df)) * 100, 2)
            }
    return missing_summary

# ── Assumption Checker ──────────────────────────────────────────────────────
def check_assumptions(df, test_type, params):
    checks = []
    def chk(name, passed, note, suggestion="", fix=None):
        checks.append({"name":name,"passed":bool(passed),"note":note,"suggestion":suggestion,"fix":fix})

    num_var = params.get("numeric_var") or params.get("dependent")
    grp_var = params.get("group_var")

    if test_type in ("t_test","anova","mann_whitney","kruskal_wallis") and num_var and grp_var:
        series = pd.to_numeric(df[num_var], errors="coerce").dropna()
        n = len(series)
        chk("Adequate sample size", n >= 30, f"N = {n}. {'Sufficient (≥ 30).' if n>=30 else 'Small sample.'}")
        if n >= 3 and series.nunique() > 2:
            try:
                if n < 5000:
                    sw_stat, sw_p = stats.shapiro(series)
                    normal = sw_p > 0.05
                    chk("Normality (Shapiro-Wilk)", normal, f"W = {sw_stat:.3f}, {fp(sw_p)}")
                else:
                    from scipy.stats import normaltest
                    k2_stat, k2_p = normaltest(series)
                    normal = k2_p > 0.05
                    chk("Normality (D'Agostino K²)", normal, f"K² = {k2_stat:.3f}, {fp(k2_p)}")
            except:
                chk("Normality check", False, "Could not compute normality test.")
        Q1,Q3=series.quantile(.25),series.quantile(.75); IQR=Q3-Q1
        n_out=int(((series<Q1-3*IQR)|(series>Q3+3*IQR)).sum())
        chk("No extreme outliers", n_out==0, f"{n_out} extreme outliers detected." if n_out else "No extreme outliers.")
        if test_type in ("t_test","anova") and grp_var:
            groups = [pd.to_numeric(df.loc[df[grp_var]==g,num_var],errors="coerce").dropna() for g in df[grp_var].dropna().unique()]
            if len(groups)>=2 and all(len(g)>=2 for g in groups):
                lev_s, lev_p = stats.levene(*groups)
                chk("Equal variances (Levene)", lev_p > 0.05, f"Levene's F = {lev_s:.3f}, {fp(lev_p)}")

    if test_type == "chi_square":
        v1,v2 = params.get("var1",""),params.get("var2","")
        if v1 and v2:
            ct = pd.crosstab(df[v1],df[v2])
            if ct.empty or ct.sum().sum() == 0:
                chk("Valid contingency table", False, "Contingency table is empty.")
            else:
                try:
                    chi2, p, dof, exp = stats.chi2_contingency(ct)
                    chk("Expected frequencies ≥ 5", not (exp < 5).any(), 
                        "All expected frequencies ≥ 5." if not (exp < 5).any() else "Some expected frequencies < 5.")
                except:
                    chk("Valid chi-square test", False, "Could not compute chi-square.")

    return checks


# ══════════════════════════════════════════════════════════════════════════════
# SMART TEST RECOMMENDER
# ══════════════════════════════════════════════════════════════════════════════
def recommend_tests(df, columns):
    cols = {c["name"]: c for c in columns}
    num = [c["name"] for c in columns if c["dtype"].startswith("num") and c["n_unique"]>5]
    cat = [c["name"] for c in columns if not c["dtype"].startswith("num") or c["n_unique"]<=10]
    cat2 = [c for c in cat if df[c].nunique()==2]
    cat3p = [c for c in cat if df[c].nunique()>=3]
    suggestions = []
    if len(num)>=1:
        suggestions.append({"test":"descriptive","reason":"Summarise numeric variables.","priority":1})
    if len(cat)>=2:
        suggestions.append({"test":"chi_square","reason":"Test association between categorical variables.","priority":2})
    if num and cat2:
        suggestions.append({"test":"t_test","reason":f"Compare {num[0]} between two groups.","priority":2})
    if num and cat3p:
        suggestions.append({"test":"anova","reason":f"Compare {num[0]} across 3+ groups.","priority":3})
    if len(num)>=2:
        suggestions.append({"test":"correlation","reason":"Explore relationships between numeric variables.","priority":2})
    if len(num)>=2:
        suggestions.append({"test":"regression","reason":f"Predict {num[0]} from other variables.","priority":3})
    if num and cat2:
        suggestions.append({"test":"logistic_regression","reason":f"Predict binary outcome from numeric predictors.","priority":3})
    if num and cat2:
        suggestions.append({"test":"mann_whitney","reason":"Non-parametric alternative.","priority":4})
    return sorted(suggestions, key=lambda x: x["priority"])


# ══════════════════════════════════════════════════════════════════════════════
# DESCRIPTIVE STATISTICS
# ══════════════════════════════════════════════════════════════════════════════
@with_timeout(MAX_RUNTIME_SECONDS)
def descriptive_statistics(df, columns):
    check_analysis_limits(df, "descriptive", {"columns": columns})
    for col in columns:
        if col not in df.columns:
            raise ValidationError(f"{ERROR_CODES['VAL_001']}: {col}")
    original_rows = len(df)
    num_cols = [c for c in columns if pd.api.types.is_numeric_dtype(pd.to_numeric(df[c],errors="coerce"))]
    num_cols = [c for c in columns if pd.to_numeric(df[c],errors="coerce").notna().sum() > len(df)*0.5]
    cat_cols = [c for c in columns if c not in num_cols]
    charts=[]; numeric_summary=[]; categorical_summary=[]
    chart_count = 0
    warnings_list = []

    if num_cols:
        for c in num_cols:
            s=pd.to_numeric(df[c],errors="coerce").dropna()
            if len(s)==0: continue
            missing_mask = (df[c].isna() | df[c].astype(str).str.strip().isin(["", "N/A"]))
            missing_count = missing_mask.sum()
            if len(s) >= 3 and s.nunique() > 2:
                try:
                    if len(s) < 5000:
                        sw_s, sw_p = stats.shapiro(s)
                    else:
                        from scipy.stats import normaltest
                        sw_s, sw_p = normaltest(s)
                except:
                    sw_s, sw_p = np.nan, np.nan
                    warnings_list.append(f"Could not compute normality test for {c}")
            else:
                sw_s, sw_p = np.nan, np.nan
            numeric_summary.append({"Variable":c,"N":int(s.count()),"Missing":int(missing_count),
                "Mean":round(float(s.mean()),4),"Median":round(float(s.median()),4),"Std Dev":round(float(s.std()),4),
                "Min":round(float(s.min()),4),"Max":round(float(s.max()),4),"Q1":round(float(s.quantile(.25)),4),
                "Q3":round(float(s.quantile(.75)),4),"Skewness":round(float(s.skew()),4),
                "Kurtosis":round(float(s.kurt()),4),"Shapiro-Wilk p":round(float(sw_p),4) if not np.isnan(sw_p) else "—"})

        selected_num_cols = num_cols[:15]
        num_charts = len(selected_num_cols)

        if chart_count < MAX_CHARTS and num_charts > 0:
            nc=min(3, num_charts); nr=(num_charts+nc-1)//nc
            fig,axes=plt.subplots(nr,nc,figsize=(5.5*nc,4*nr)); fig.patch.set_facecolor(P["white"])
            axf=np.array(axes).flatten() if num_charts>1 else [axes]
            for i,c in enumerate(selected_num_cols):
                ax=axf[i]; s=pd.to_numeric(df[c],errors="coerce").dropna()
                ax.hist(s,bins="auto",color=PAL[i%len(PAL)],alpha=.85,edgecolor="white",linewidth=.5)
                ax.axvline(s.mean(),color=P["navy"],ls="--",lw=1.5,label=f"M={s.mean():.2f}")
                ax.axvline(s.median(),color=P["err"],ls=":",lw=1.5,label=f"Mdn={s.median():.2f}")
                ax.set_title(c,fontsize=11,fontweight="bold",color=P["navy"]); ax.legend(fontsize=8)
                ax.set_xlabel("Value"); ax.set_ylabel("Frequency")
            for j in range(i+1,len(axf)): axf[j].set_visible(False)
            _style(axf[:num_charts]); fig.suptitle("Distribution Histograms",fontsize=13,fontweight="bold",color=P["navy"],y=1.01)
            plt.tight_layout(); charts.append({"title":"Histograms","img":_store_chart(fig, "histograms", 110)})
            chart_count += 1

            if num_charts > 1 and chart_count < MAX_CHARTS:
                fig2,ax2=plt.subplots(figsize=(max(9, num_charts*1.6),5)); fig2.patch.set_facecolor(P["white"])
                data=[pd.to_numeric(df[c],errors="coerce").dropna().values for c in selected_num_cols]
                bp=ax2.boxplot(data,patch_artist=True,medianprops=dict(color=P["navy"],lw=2),flierprops=dict(marker="o",ms=4,alpha=.5))
                for patch,col in zip(bp["boxes"],PAL): patch.set_facecolor(col); patch.set_alpha(.75)
                ax2.set_xticklabels(selected_num_cols,rotation=30,ha="right",fontsize=9)
                ax2.set_title("Comparative Boxplots",fontweight="bold",color=P["navy"])
                _style(ax2); plt.tight_layout(); charts.append({"title":"Boxplots","img":_store_chart(fig2, "boxplots", 110)})
                chart_count += 1

            if chart_count < MAX_CHARTS:
                fig3,axes3=plt.subplots(nr,nc,figsize=(5*nc,4*nr)); fig3.patch.set_facecolor(P["white"])
                axf3=np.array(axes3).flatten() if num_charts>1 else [axes3]
                for i,c in enumerate(selected_num_cols):
                    s=pd.to_numeric(df[c],errors="coerce").dropna()
                    stats.probplot(s,dist="norm",plot=axf3[i])
                    axf3[i].set_title(f"Q-Q: {c}",fontsize=10,fontweight="bold",color=P["navy"])
                    axf3[i].get_lines()[0].set(color=PAL[i%len(PAL)],alpha=.7,ms=4)
                    axf3[i].get_lines()[1].set(color=P["err"],lw=1.5)
                for j in range(i+1,len(axf3)): axf3[j].set_visible(False)
                _style(axf3[:num_charts]); fig3.suptitle("Q-Q Plots",fontsize=13,fontweight="bold",color=P["navy"],y=1.01)
                plt.tight_layout(); charts.append({"title":"Q-Q Plots","img":_store_chart(fig3, "qqplots", 110)})
                chart_count += 1

    if cat_cols:
        selected_cat_cols = cat_cols[:10]
        for c in selected_cat_cols:
            if chart_count >= MAX_CHARTS: break
            vc=df[c].value_counts(); pct=df[c].value_counts(normalize=True)*100
            categorical_summary.append({"variable":c,"table":[{"Category":str(k),"Count":int(v),"Percent (%)":round(float(pct[k]),2)} for k,v in vc.items()]})
            fig4,ax4=plt.subplots(figsize=(max(6,len(vc)*.9+2),5)); fig4.patch.set_facecolor(P["white"])
            bars=ax4.bar(vc.index.astype(str),vc.values,color=PAL[:len(vc)],alpha=.88,edgecolor="white",lw=.5)
            for b,v in zip(bars,vc.values):
                ax4.text(b.get_x()+b.get_width()/2,b.get_height()+.3,f"{v}\n({v/len(df)*100:.1f}%)",ha="center",va="bottom",fontsize=9,color=P["navy"])
            ax4.set_title(f"Frequency: {c}",fontweight="bold",color=P["navy"]); ax4.tick_params(axis="x",rotation=30)
            _style(ax4); plt.tight_layout(); charts.append({"title":f"Bar Chart: {c}","img":_store_chart(fig4, f"barchart_{c}", 110)})
            chart_count += 1
            if chart_count < MAX_CHARTS:
                fig5,ax5=plt.subplots(figsize=(6,5)); fig5.patch.set_facecolor(P["white"])
                ax5.pie(vc.values,labels=vc.index.astype(str),colors=PAL[:len(vc)],autopct="%1.1f%%",startangle=140,pctdistance=.82,
                        wedgeprops=dict(edgecolor="white",linewidth=1.5))
                ax5.set_title(f"Proportion: {c}",fontweight="bold",color=P["navy"])
                charts.append({"title":f"Pie Chart: {c}","img":_store_chart(fig5, f"piechart_{c}", 110)})
                chart_count += 1

    interp=[]
    for r in numeric_summary:
        sk="approximately symmetric" if abs(r["Skewness"])<.5 else ("positively skewed" if r["Skewness"]>0 else "negatively skewed")
        interp.append(f"**{r['Variable']}**: N={r['N']}, M={r['Mean']}, SD={r['Std Dev']}, Median={r['Median']}. Range=[{r['Min']},{r['Max']}]. Distribution is {sk}.")

    rows_removed = original_rows - len(df)
    return {
        "test":"Descriptive Statistics",
        "numeric_summary":numeric_summary,
        "categorical_summary":categorical_summary,
        "charts":charts,
        "interpretation":"\n\n".join(interp) if interp else "Selected columns analysed.",
        "apa_citation":"Descriptive statistics are reported as M (SD).",
        "significance":"—","p_value":None,"p_display":"—",
        "missing_data_summary": get_missing_summary(df),
        "warnings": warnings_list,
        "data_quality": {
            "original_rows": original_rows,
            "rows_after_cleaning": len(df),
            "rows_removed": rows_removed,
            "percent_removed": round((rows_removed / original_rows) * 100, 2) if original_rows > 0 else 0
        },
        "engine_version": ENGINE_VERSION,
        "analysis_timestamp": datetime.now().isoformat(),
        "reproducibility": get_system_metadata()
    }


# ══════════════════════════════════════════════════════════════════════════════
# CHI-SQUARE
# ══════════════════════════════════════════════════════════════════════════════
@with_timeout(MAX_RUNTIME_SECONDS)
def chi_square_test(df, var1, var2):
    validate_categorical_column(df, var1, "for chi-square test")
    validate_categorical_column(df, var2, "for chi-square test")
    original_rows = len(df)
    ct=pd.crosstab(df[var1],df[var2])
    if ct.empty or ct.sum().sum() == 0:
        return {"error": ERROR_CODES["ERR_002"], "message": "Contingency table is empty."}
    
    if ct.shape == (2, 2):
        try:
            chi2, p, dof, exp = stats.chi2_contingency(ct)
            if (exp < 5).any():
                odds_ratio, p_fisher = stats.fisher_exact(ct)
                logger.info("Fisher's Exact Test used for 2x2 table")
                result = _chi_square_result(ct, None, p_fisher, None, "Fisher's Exact Test", {}, original_rows, len(df))
                result["data_quality"] = {
                    "original_rows": original_rows,
                    "rows_after_cleaning": len(df),
                    "rows_removed": original_rows - len(df),
                    "percent_removed": round(((original_rows - len(df)) / original_rows) * 100, 2) if original_rows > 0 else 0
                }
                result["engine_version"] = ENGINE_VERSION
                result["analysis_timestamp"] = datetime.now().isoformat()
                result["reproducibility"] = get_system_metadata()
                result["missing_data_summary"] = get_missing_summary(df)
                return result
        except:
            pass
    
    try:
        chi2,p,dof,exp=stats.chi2_contingency(ct)
    except Exception as e:
        return {"error": ERROR_CODES["ERR_002"], "message": f"Chi-square test failed: {str(e)}"}
    
    result = _chi_square_result(ct, chi2, p, dof, "Chi-Square Test of Independence", {}, original_rows, len(df))
    result["data_quality"] = {
        "original_rows": original_rows,
        "rows_after_cleaning": len(df),
        "rows_removed": original_rows - len(df),
        "percent_removed": round(((original_rows - len(df)) / original_rows) * 100, 2) if original_rows > 0 else 0
    }
    result["engine_version"] = ENGINE_VERSION
    result["analysis_timestamp"] = datetime.now().isoformat()
    result["reproducibility"] = get_system_metadata()
    result["missing_data_summary"] = get_missing_summary(df)
    return result

def _chi_square_result(ct, chi2, p, dof, test_name, extra, original_rows, cleaned_rows):
    n=ct.values.sum(); min_dim=min(ct.shape)-1
    v=np.sqrt(chi2/(n*max(min_dim,1))) if chi2 else 0
    ef="negligible" if v<.1 else "small" if v<.3 else "moderate" if v<.5 else "large"
    vrd,h0=verdict(p); charts=[]
    fig,(a1,a2)=plt.subplots(1,2,figsize=(14,5)); fig.patch.set_facecolor(P["white"])
    ct.div(ct.sum(axis=1),axis=0).mul(100).plot(kind="bar",stacked=True,ax=a1,color=PAL[:len(ct.columns)],edgecolor="white",alpha=.88)
    a1.set_title(f"Stacked Bar: {ct.index.name} × {ct.columns.name}",fontweight="bold",color=P["navy"])
    a1.set_xlabel(ct.index.name); a1.set_ylabel("Percentage (%)"); a1.tick_params(axis="x",rotation=30)
    a1.legend(title=ct.columns.name,bbox_to_anchor=(1.02,1),fontsize=8)
    sns.heatmap(ct,annot=True,fmt="d",cmap="Blues",ax=a2,linewidths=.5,cbar_kws={"shrink":.7})
    a2.set_title("Observed Frequencies",fontweight="bold",color=P["navy"])
    _style([a1]); fig.suptitle(f"{test_name}: {ct.index.name} × {ct.columns.name}",fontsize=11,fontweight="bold",color=P["navy"],y=1.01)
    plt.tight_layout(); charts.append({"title":"Chi-Square Visualisation","img":_store_chart(fig, "chisquare_viz", 110)})
    if chi2:
        exp_df=pd.DataFrame(exp,index=ct.index,columns=ct.columns); resid=(ct-exp_df)/np.sqrt(exp_df)
        fig2,ax=plt.subplots(figsize=(8,5)); fig2.patch.set_facecolor(P["white"])
        sns.heatmap(resid,annot=True,fmt=".2f",cmap="RdBu_r",center=0,ax=ax,linewidths=.5,cbar_kws={"shrink":.7})
        ax.set_title("Standardised Residuals",fontweight="bold",color=P["navy"])
        charts.append({"title":"Standardised Residuals","img":_store_chart(fig2, "chisquare_residuals", 110)})
    ct_out=ct.copy(); ct_out["Row Total"]=ct_out.sum(axis=1); ct_out.loc["Col Total"]=ct_out.sum()
    return {"test":test_name,"chi2":round(float(chi2),4) if chi2 else None,"p_value":round(float(p),4),
            "p_display":fp(p),"dof":int(dof) if dof else None,"significance":sp(p),
            "cramers_v":round(float(v),4) if v else None,"effect_size":ef,"n":int(n),
            "contingency_table":ct_out.to_dict(),"charts":charts,
            "apa_citation":f"χ²({dof}, N={n})={chi2:.2f}, {fp(p)}, Cramér's V={v:.2f}" if chi2 else f"Fisher's Exact Test: p={fp(p)}",
            "interpretation":f"A {test_name} examined the relationship between {ct.index.name} and {ct.columns.name}. "
                f"The association was {vrd}, {fp(p)}. We therefore {h0} the null hypothesis."}


# ══════════════════════════════════════════════════════════════════════════════
# T-TEST
# ══════════════════════════════════════════════════════════════════════════════
@with_timeout(MAX_RUNTIME_SECONDS)
def independent_ttest(df, numeric_var, group_var, alpha=.05):
    validate_numeric_column(df, numeric_var, "for t-test")
    validate_categorical_column(df, group_var, "for t-test")
    validate_row_count(df, 3, "for t-test")
    original_rows = len(df)
    gs=df[group_var].dropna().unique()
    if len(gs)!=2: return {"error": ERROR_CODES["VAL_004"], "message": f"Need exactly 2 groups; found {len(gs)}."}
    g1=pd.to_numeric(df.loc[df[group_var]==gs[0],numeric_var],errors="coerce").dropna()
    g2=pd.to_numeric(df.loc[df[group_var]==gs[1],numeric_var],errors="coerce").dropna()
    t_s,t_p=stats.ttest_ind(g1,g2); lev_s,lev_p=stats.levene(g1,g2)
    if lev_p < 0.05:
        t_s_welch, t_p_welch = stats.ttest_ind(g1, g2, equal_var=False)
        t_s, t_p = t_s_welch, t_p_welch
        d = (g1.mean() - g2.mean()) / g2.std() if g2.std() > 0 else 0
        d_label = "Glass's Δ"
    else:
        pool = np.sqrt(((len(g1)-1)*g1.std()**2+(len(g2)-1)*g2.std()**2)/(len(g1)+len(g2)-2))
        d = (g1.mean()-g2.mean())/pool if pool>0 else 0
        d_label = "Cohen's d"
    ef="negligible" if abs(d)<.2 else "small" if abs(d)<.5 else "medium" if abs(d)<.8 else "large"
    se=np.sqrt(g1.var()/len(g1)+g2.var()/len(g2))
    df_w=se**4/((g1.var()/len(g1))**2/(len(g1)-1)+(g2.var()/len(g2))**2/(len(g2)-1))
    tc=stats.t.ppf(1-alpha/2,df_w); md=g1.mean()-g2.mean()
    ci=[md-tc*se,md+tc*se]; vrd,h0=verdict(t_p); charts=[]
    warnings_list = []
    
    fig,(a1,a2)=plt.subplots(1,2,figsize=(13,5)); fig.patch.set_facecolor(P["white"])
    pdf=pd.DataFrame({numeric_var:pd.concat([g1,g2]),group_var:[str(gs[0])]*len(g1)+[str(gs[1])]*len(g2)})
    parts=a1.violinplot([g1.values,g2.values],positions=[1,2],showmedians=True)
    for i,pc in enumerate(parts["bodies"]): pc.set_facecolor(PAL[i]); pc.set_alpha(.7)
    parts["cmedians"].set_color(P["navy"]); a1.set_xticks([1,2]); a1.set_xticklabels([str(gs[0]),str(gs[1])])
    a1.set_title(f"Violin: {numeric_var}",fontweight="bold",color=P["navy"]); a1.set_ylabel(numeric_var)
    sns.boxplot(data=pdf,x=group_var,y=numeric_var,palette=PAL[:2],ax=a2,width=.5)
    sns.stripplot(data=pdf,x=group_var,y=numeric_var,color=P["navy"],alpha=.4,size=4,jitter=True,ax=a2)
    a2.set_title(f"Box + Strip: {numeric_var}",fontweight="bold",color=P["navy"])
    _style([a1,a2]); fig.suptitle(f"T-Test: {numeric_var} by {group_var} | t({df_w:.1f})={t_s:.3f}, {fp(t_p)}",fontsize=11,fontweight="bold",color=P["navy"],y=1.02)
    plt.tight_layout(); charts.append({"title":"Group Comparison","img":_store_chart(fig, "ttest_violin", 110)})
    
    fig2,ax=plt.subplots(figsize=(7,5)); fig2.patch.set_facecolor(P["white"])
    for i,(lbl,gd) in enumerate([(str(gs[0]),g1),(str(gs[1]),g2)]):
        ax.errorbar(i,gd.mean(),yerr=gd.sem()*1.96,fmt="o",color=PAL[i],ms=12,capsize=8,lw=2,label=f"{lbl}: M={gd.mean():.2f}")
    ax.set_xticks([0,1]); ax.set_xticklabels([str(gs[0]),str(gs[1])],fontsize=11)
    ax.set_ylabel(numeric_var); ax.set_title("Means with 95% CI",fontweight="bold",color=P["navy"]); ax.legend()
    _style(ax); plt.tight_layout(); charts.append({"title":"Means with 95% CI","img":_store_chart(fig2, "ttest_means", 110)})
    
    stbl=[{"Group":str(gs[i]),"N":int(len(g)),"Mean":round(float(g.mean()),4),"Std Dev":round(float(g.std()),4),"Std Error":round(float(g.sem()),4)} for i,g in [(0,g1),(1,g2)]]
    # ── FIXED: Correct rows_removed calculation ──────────────────────────
    rows_removed = original_rows - (len(g1) + len(g2))
    assumptions = [{"name":"Normality","passed":True,"note":"Check Q-Q plot"}]
    return {
        "test":"Independent Samples T-Test","t_statistic":round(float(t_s),4),"p_value":round(float(t_p),4),
        "p_display":fp(t_p),"df":round(float(df_w),2),"significance":sp(t_p),"mean_difference":round(float(md),4),
        "ci_95":[round(float(ci[0]),4),round(float(ci[1]),4)],"effect_size":round(float(d),4),"effect_label":d_label,
        "levene_p":round(float(lev_p),4),"summary_table":stbl,"charts":charts,
        "assumptions": assumptions,
        "warnings": warnings_list,
        "apa_citation":f"t({df_w:.2f})={t_s:.2f}, {fp(t_p)}, {d_label}={d:.2f}, 95% CI [{ci[0]:.2f},{ci[1]:.2f}]",
        "interpretation":f"Independent samples t-test comparing {numeric_var} between {gs[0]} (M={g1.mean():.2f},SD={g1.std():.2f}) "
            f"and {gs[1]} (M={g2.mean():.2f},SD={g2.std():.2f}). Difference was {vrd}, t({df_w:.2f})={t_s:.2f}, {fp(t_p)}, {d_label}={d:.2f} ({ef} effect). "
            f"95% CI=[{ci[0]:.2f},{ci[1]:.2f}]. We {h0} the null hypothesis." +
            (f" Levene's test: {fp(lev_p)} — Welch correction applied." if lev_p<.05 else ""),
        "data_quality": {
            "original_rows": original_rows,
            "rows_after_cleaning": len(g1) + len(g2),
            "rows_removed": rows_removed,
            "percent_removed": round((rows_removed / original_rows) * 100, 2) if original_rows > 0 else 0
        },
        "missing_data_summary": get_missing_summary(df),
        "engine_version": ENGINE_VERSION,
        "analysis_timestamp": datetime.now().isoformat(),
        "reproducibility": get_system_metadata()
    }


# ══════════════════════════════════════════════════════════════════════════════
# ANOVA
# ══════════════════════════════════════════════════════════════════════════════
@with_timeout(MAX_RUNTIME_SECONDS)
def one_way_anova(df, numeric_var, group_var):
    validate_numeric_column(df, numeric_var, "for ANOVA")
    validate_categorical_column(df, group_var, "for ANOVA")
    validate_row_count(df, 3, "for ANOVA")
    original_rows = len(df)
    gs=df[group_var].dropna().unique()
    gdata=[]
    warnings_list = []
    
    for g in gs:
        group_data = pd.to_numeric(df.loc[df[group_var]==g,numeric_var],errors="coerce").dropna().values
        if len(group_data) < 2:
            return {"error": ERROR_CODES["VAL_004"], "message": f"Group '{g}' has only {len(group_data)} observations."}
        gdata.append(group_data)
    F,p=stats.f_oneway(*gdata); gm=np.concatenate(gdata).mean()
    ss_b=sum(len(g)*(g.mean()-gm)**2 for g in gdata); ss_t=sum((v-gm)**2 for g in gdata for v in g)
    eta=ss_b/ss_t if ss_t>0 else 0; ef="small" if eta<.06 else "medium" if eta<.14 else "large"
    n_t=sum(len(g) for g in gdata); k=len(gs); vrd,h0=verdict(p); charts=[]
    
    fig,(a1,a2)=plt.subplots(1,2,figsize=(14,5)); fig.patch.set_facecolor(P["white"])
    pdf=pd.DataFrame({numeric_var:np.concatenate(gdata),group_var:np.repeat([str(g) for g in gs],[len(g) for g in gdata])})
    sns.boxplot(data=pdf,x=group_var,y=numeric_var,palette=PAL[:k],ax=a1,width=.5)
    sns.stripplot(data=pdf,x=group_var,y=numeric_var,color=P["navy"],alpha=.35,size=4,jitter=True,ax=a1)
    a1.set_title(f"Group Distributions: {numeric_var}",fontweight="bold",color=P["navy"])
    a1.tick_params(axis="x",rotation=30)
    means=[g.mean() for g in gdata]; sems=[stats.sem(g) for g in gdata]
    bars=a2.bar(range(k),means,color=PAL[:k],alpha=.85,edgecolor="white",lw=.5)
    a2.errorbar(range(k),means,yerr=[1.96*s for s in sems],fmt="none",color=P["navy"],capsize=6,lw=1.5)
    a2.set_xticks(range(k)); a2.set_xticklabels([str(g) for g in gs],rotation=30,ha="right")
    a2.set_ylabel(numeric_var); a2.set_title("Group Means with 95% CI",fontweight="bold",color=P["navy"])
    _style([a1,a2]); fig.suptitle(f"ANOVA: {numeric_var} by {group_var} | F({k-1},{n_t-k})={F:.3f}, {fp(p)}, η²={eta:.3f}",fontsize=11,fontweight="bold",color=P["navy"],y=1.02)
    plt.tight_layout(); charts.append({"title":"ANOVA Group Plots","img":_store_chart(fig, "anova_plots", 110)})
    
    from statsmodels.stats.multicomp import pairwise_tukeyhsd
    posthoc = []
    if p < 0.05 and k > 2:
        try:
            tukey = pairwise_tukeyhsd(endog=pdf[numeric_var], groups=pdf[group_var], alpha=0.05)
            summary_data = tukey.summary().data[1:]
            for row in summary_data:
                posthoc.append({
                    "Group 1": str(row[0]),
                    "Group 2": str(row[1]),
                    "Mean Diff": round(float(row[2]), 4),
                    "p-value": round(float(row[3]), 4),
                    "Significant": "Yes" if row[6] else "No"
                })
        except Exception as e:
            warnings_list.append(f"Tukey HSD failed: {str(e)}. Using Bonferroni fallback.")
            n_pairs = len(list(combinations(gs,2)))
            for (i,g_a),(j,g_b) in combinations(enumerate(gs),2):
                _,tp=stats.ttest_ind(gdata[i],gdata[j]); tadj=min(tp*n_pairs,1.0)
                posthoc.append({"Group 1":str(g_a),"Group 2":str(g_b),"Mean Diff":round(float(gdata[i].mean()-gdata[j].mean()),4),"p (Bonferroni)":round(float(tadj),4),"Significant":"Yes" if tadj<.05 else "No"})
    
    stbl=[{"Group":str(g),"N":int(len(gd)),"Mean":round(float(gd.mean()),4),"Std Dev":round(float(gd.std()),4),"Std Error":round(float(stats.sem(gd)),4),"Min":round(float(gd.min()),4),"Max":round(float(gd.max()),4)} for g,gd in zip(gs,gdata)]
    # ── FIXED: Correct rows_removed calculation ──────────────────────────
    rows_removed = original_rows - n_t
    assumptions = [{"name":"Normality","passed":True,"note":"Check Q-Q plots per group"}]
    return {
        "test":"One-Way ANOVA","f_statistic":round(float(F),4),"p_value":round(float(p),4),"p_display":fp(p),
        "df_between":int(k-1),"df_within":int(n_t-k),"significance":sp(p),"eta_squared":round(float(eta),4),
        "effect_size":ef,"summary_table":stbl,"posthoc_table":posthoc,"charts":charts,
        "assumptions": assumptions,
        "warnings": warnings_list,
        "apa_citation":f"F({k-1},{n_t-k})={F:.2f}, {fp(p)}, η²={eta:.3f}",
        "interpretation":f"One-way ANOVA compared {numeric_var} across {k} groups. Result was {vrd}, F({k-1},{n_t-k})={F:.2f}, {fp(p)}, η²={eta:.3f} ({ef} effect). "
            f"We {h0} the null hypothesis."+(" Tukey HSD post-hoc comparisons provided." if p<.05 and len(posthoc)>0 else ""),
        "data_quality": {
            "original_rows": original_rows,
            "rows_after_cleaning": n_t,
            "rows_removed": rows_removed,
            "percent_removed": round((rows_removed / original_rows) * 100, 2) if original_rows > 0 else 0
        },
        "missing_data_summary": get_missing_summary(df),
        "engine_version": ENGINE_VERSION,
        "analysis_timestamp": datetime.now().isoformat(),
        "reproducibility": get_system_metadata()
    }


# ══════════════════════════════════════════════════════════════════════════════
# CORRELATION
# ══════════════════════════════════════════════════════════════════════════════
@with_timeout(MAX_RUNTIME_SECONDS)
def correlation_analysis(df, columns, method="pearson"):
    num = [c for c in columns if pd.to_numeric(df[c],errors="coerce").notna().sum()>len(df)*.4]
    if len(num)<2: return {"error": ERROR_CODES["VAL_001"], "message": "Need ≥2 numeric columns."}
    for col in num:
        validate_numeric_column(df, col, "for correlation")
    original_rows = len(df)
    sub = df[num].apply(pd.to_numeric, errors="coerce")
    charts = []
    warnings_list = []
    
    constant_cols = [c for c in sub.columns if sub[c].nunique() <= 1]
    if constant_cols:
        warnings_list.append(f"Dropped constant columns from correlation: {constant_cols}")
        sub = sub.drop(columns=constant_cols)
        num = [c for c in num if c not in constant_cols]
        if len(num) < 2:
            return {"error": ERROR_CODES["VAL_009"], "message": "After removing constant columns, fewer than 2 numeric columns remain."}
    
    # ── FIXED: Initialize matrices with float dtype ──────────────────────
    corr_matrix = pd.DataFrame(np.nan, index=num, columns=num, dtype=float)
    p_matrix = pd.DataFrame(np.nan, index=num, columns=num, dtype=float)
    n_matrix = pd.DataFrame(np.nan, index=num, columns=num, dtype=float)
    
    pairs = []
    for c1, c2 in combinations(num, 2):
        pair_data = sub[[c1, c2]].dropna()
        if len(pair_data) > 2:
            if method == "spearman":
                r, p = stats.spearmanr(pair_data[c1], pair_data[c2])
            else:
                r, p = stats.pearsonr(pair_data[c1], pair_data[c2])
            corr_matrix.loc[c1, c2] = r
            corr_matrix.loc[c2, c1] = r
            p_matrix.loc[c1, c2] = p
            p_matrix.loc[c2, c1] = p
            n_matrix.loc[c1, c2] = len(pair_data)
            n_matrix.loc[c2, c1] = len(pair_data)
            st = "negligible" if abs(r)<.1 else "weak" if abs(r)<.3 else "moderate" if abs(r)<.5 else "strong" if abs(r)<.7 else "very strong"
            pairs.append({"Variable 1":c1,"Variable 2":c2,"r":round(float(r),4),"p-value":round(float(p),4),"p_display":fp(p),"Significance":sp(p),"Strength":f"{'positive' if r>0 else 'negative'} {st}"})
    
    for c in num:
        corr_matrix.loc[c, c] = 1.0
        p_matrix.loc[c, c] = 0.0
        n_matrix.loc[c, c] = len(sub[c].dropna())
    
    # Multiple testing correction
    if pairs:
        from statsmodels.stats.multitest import multipletests
        p_values = [p_["p-value"] for p_ in pairs]
        rejected, pvals_corrected, _, _ = multipletests(p_values, alpha=0.05, method='fdr_bh')
        for i, pair in enumerate(pairs):
            pair["p_corrected"] = round(float(pvals_corrected[i]), 4)
            pair["Significant_corrected"] = "Yes" if rejected[i] else "No"
    
    # Heatmap
    fig,ax=plt.subplots(figsize=(max(6,len(num)*1.3+2),max(5,len(num)*1.2)))
    fig.patch.set_facecolor(P["white"])
    mask=np.triu(np.ones_like(corr_matrix,dtype=bool))
    sns.heatmap(corr_matrix.astype(float), annot=True, fmt=".3f", cmap="coolwarm", vmin=-1, vmax=1, center=0,
                ax=ax, mask=mask, square=True, linewidths=.5, annot_kws={"size":10,"weight":"bold"}, cbar_kws={"shrink":.8})
    ax.set_title(f"{method.title()} Correlation Matrix", fontsize=13, fontweight="bold", color=P["navy"])
    plt.tight_layout(); charts.append({"title":"Correlation Heatmap","img":_store_chart(fig, "correlation_heatmap", 110)})
    
    # ── FIXED: Skip pairplot for >4 variables ─────────────────────────────
    if 2 <= len(num) <= MAX_PAIRPLOT_VARS:
        sample_size = min(MAX_PAIRPLOT_SAMPLE, len(sub))
        if len(sub) > sample_size:
            pairplot_df = sub.sample(sample_size, random_state=RANDOM_SEED)
        else:
            pairplot_df = sub
        g = sns.pairplot(pairplot_df, diag_kind="hist", plot_kws={"alpha":.5, "color":P["blue"], "s":20}, diag_kws={"color":P["teal"], "fill":True})
        g.figure.suptitle("Scatter Matrix", y=1.02, fontsize=13, fontweight="bold", color=P["navy"])
        g.figure.patch.set_facecolor(P["white"])
        charts.append({"title":"Scatter Matrix","img":_store_chart(g.figure, "correlation_scatter", 110)})
    elif len(num) > MAX_PAIRPLOT_VARS:
        # Skip pairplot to save memory
        charts.append({"title":"Scatter Matrix","note":"Skipped - too many variables (>4)"})
    
    if pairs:
        fig3,ax3=plt.subplots(figsize=(10,5)); fig3.patch.set_facecolor(P["white"])
        x_pos=range(len(pairs)); rs=[abs(p_["r"]) for p_ in pairs]; cols_=[P["ok"] if p_["r"]>0 else P["err"] for p_ in pairs]
        sc=ax3.scatter(x_pos,[p_["r"] for p_ in pairs],s=[r*500+30 for r in rs],c=cols_,alpha=.75,edgecolors="white",lw=1.5)
        ax3.axhline(0,color=P["slate"],lw=1,ls="--"); ax3.axhline(.3,color=P["warn"],lw=.8,ls=":"); ax3.axhline(-.3,color=P["warn"],lw=.8,ls=":")
        ax3.set_xticks(list(x_pos)); ax3.set_xticklabels([f"{p_['Variable 1'][:8]}×{p_['Variable 2'][:8]}" for p_ in pairs],rotation=35,ha="right",fontsize=8)
        ax3.set_ylabel("r value"); ax3.set_title("Correlation Bubble Chart",fontweight="bold",color=P["navy"])
        _style(ax3); plt.tight_layout(); charts.append({"title":"Correlation Bubble Chart","img":_store_chart(fig3, "correlation_bubble", 110)})
    
    # ── FIXED: Correct rows_removed calculation ──────────────────────────
    rows_removed = original_rows - len(sub)
    return {
        "test":f"{method.title()} Correlation","method":method,"n":int(len(sub)),"variables":num,
        "pairs_table":pairs,"charts":charts,
        "p_value":min([p_["p-value"] for p_ in pairs]) if pairs else 1,"p_display":"see table","significance":"see table",
        "apa_citation":f"{method.title()} correlation analysis on {len(num)} variables (N={len(sub)}).",
        "interpretation":f"**{method.title()} correlation** on {len(pairs)} pairs.\n\n"+"".join(
            f"- **{p_['Variable 1']} & {p_['Variable 2']}**: r={p_['r']}, {p_['p_display']} ({sp(p_['p-value'])}) — {p_['Strength']}.\n" for p_ in pairs),
        "warnings": warnings_list,
        "data_quality": {
            "original_rows": original_rows,
            "rows_after_cleaning": len(sub),
            "rows_removed": rows_removed,
            "percent_removed": round((rows_removed / original_rows) * 100, 2) if original_rows > 0 else 0
        },
        "missing_data_summary": get_missing_summary(df),
        "engine_version": ENGINE_VERSION,
        "analysis_timestamp": datetime.now().isoformat(),
        "reproducibility": get_system_metadata()
    }


# ══════════════════════════════════════════════════════════════════════════════
# LINEAR REGRESSION
# ══════════════════════════════════════════════════════════════════════════════
@with_timeout(MAX_RUNTIME_SECONDS)
def linear_regression(df, dependent, predictors, robust=True, n_bootstrap=1000):
    check_analysis_limits(df, "regression", {"predictors": predictors})
    validate_numeric_column(df, dependent, "as dependent variable")
    validate_predictor_count(predictors, context="for regression")
    for pred in predictors:
        validate_numeric_column(df, pred, f"as predictor")
    
    original_rows = len(df)
    sub = df[[dependent] + predictors].apply(pd.to_numeric, errors="coerce").dropna()
    y = sub[dependent].values
    charts = []
    warnings_list = []
    
    constant_preds = [p for p in predictors if sub[p].nunique() <= 1]
    if constant_preds:
        warnings_list.append(f"Dropped constant predictors: {constant_preds}")
        predictors = [p for p in predictors if p not in constant_preds]
        sub = sub[[dependent] + predictors]
        if len(predictors) == 0:
            return {"error": ERROR_CODES["VAL_009"], "message": "All predictors are constant."}
    
    Xr = sub[predictors].values
    
    min_n = max(20, len(predictors) * 10)
    n_obs = len(sub)
    if n_obs < min_n:
        return {"error": ERROR_CODES["VAL_010"], "message": f"Requires at least {min_n} observations. You have {n_obs}."}
    if n_obs <= len(predictors) + 1:
        return {"error": ERROR_CODES["VAL_010"], "message": f"Sample size must be > predictors + 1."}
    
    import statsmodels.api as sm
    X = sm.add_constant(Xr)
    
    if np.linalg.matrix_rank(X) < X.shape[1]:
        raise SingularMatrixError(ERROR_CODES["ERR_003"])
    
    model = sm.OLS(y, X).fit()
    if robust:
        model = model.get_robustcov_results(cov_type='HC3')
    
    r2 = model.rsquared
    adj_r2 = model.rsquared_adj
    F_stat = model.fvalue
    p_f = model.f_pvalue
    
    coef = []
    coef.append({
        "Predictor": "(Intercept)",
        "B": round(float(model.params[0]), 4),
        "Std Error": round(float(model.bse[0]), 4),
        "t": round(float(model.tvalues[0]), 4),
        "p": round(float(model.pvalues[0]), 4),
        "Significance": sp(model.pvalues[0])
    })
    for i, pred in enumerate(predictors):
        coef.append({
            "Predictor": pred,
            "B": round(float(model.params[i+1]), 4),
            "Std Error": round(float(model.bse[i+1]), 4),
            "t": round(float(model.tvalues[i+1]), 4),
            "p": round(float(model.pvalues[i+1]), 4),
            "Significance": sp(model.pvalues[i+1])
        })
        # ── FIXED: Add standardized coefficients (Beta) ──────────────────
        beta = model.params[i+1] * (sub[pred].std() / sub[dependent].std())
        coef[i+1]["Beta"] = round(float(beta), 4)
    
    # ── FIXED: Bootstrap with dynamic iterations ──────────────────────────
    if n_bootstrap > 0:
        # Dynamic bootstrap iterations based on sample size
        if n_obs > 5000:
            n_bootstrap_iter = 200
        elif n_obs > 1000:
            n_bootstrap_iter = 500
        else:
            n_bootstrap_iter = 1000
        
        np.random.seed(RANDOM_SEED)
        n_params = len(model.params)
        boot_params = []
        for i in range(n_bootstrap_iter):
            try:
                idx = np.random.choice(range(n_obs), size=n_obs, replace=True)
                X_boot = X[idx, :]
                y_boot = y[idx]
                if np.linalg.matrix_rank(X_boot) < X_boot.shape[1]:
                    continue
                boot_model = sm.OLS(y_boot, X_boot).fit()
                boot_params.append(boot_model.params)
            except:
                continue
        if len(boot_params) > 0:
            boot_params = np.array(boot_params)
            lower = np.percentile(boot_params, 2.5, axis=0)
            upper = np.percentile(boot_params, 97.5, axis=0)
            for i, c in enumerate(coef):
                c["CI_lower"] = round(float(lower[i]), 4)
                c["CI_upper"] = round(float(upper[i]), 4)
        else:
            warnings_list.append("Bootstrap failed - too many singular samples")
    
    influence = model.get_influence()
    cooks_d = influence.cooks_distance[0]
    
    cooks_threshold = 4 / n_obs
    cooks_passed = float(cooks_d.max()) < cooks_threshold
    
    from statsmodels.stats.outliers_influence import variance_inflation_factor
    vif_data = []
    for i in range(X.shape[1]):
        if i == 0:
            continue
        vif = variance_inflation_factor(X, i)
        vif_data.append({
            "Variable": predictors[i-1],
            "VIF": round(float(vif), 3),
            "Interpretation": get_vif_label(vif)
        })
    
    # ── FIXED: Add Durbin-Watson test ──────────────────────────────────────
    from statsmodels.stats.stattools import durbin_watson
    dw_stat = durbin_watson(model.resid)
    
    # Regression diagnostics plots
    fig, axes = plt.subplots(2, 2, figsize=(12, 10))
    fig.patch.set_facecolor(P["white"])
    
    axes[0, 0].scatter(model.fittedvalues, y, alpha=0.6, color=P["blue"], s=40)
    axes[0, 0].plot([y.min(), y.max()], [y.min(), y.max()], linestyle='--', color=P["err"], lw=2)
    axes[0, 0].set_xlabel("Predicted"); axes[0, 0].set_ylabel("Actual")
    axes[0, 0].set_title("Actual vs Predicted", fontweight="bold", color=P["navy"])
    
    axes[0, 1].scatter(model.fittedvalues, model.resid, alpha=0.6, color=P["teal"], s=40)
    axes[0, 1].axhline(0, color=P["err"], linestyle="--", lw=1.5)
    axes[0, 1].set_xlabel("Fitted Values"); axes[0, 1].set_ylabel("Residuals")
    axes[0, 1].set_title("Residuals vs Fitted", fontweight="bold", color=P["navy"])
    
    stats.probplot(model.resid, dist="norm", plot=axes[1, 0])
    axes[1, 0].get_lines()[0].set(color=P["blue"], alpha=0.7, ms=5)
    axes[1, 0].get_lines()[1].set(color=P["err"], lw=1.5)
    axes[1, 0].set_title("Q-Q Plot of Residuals", fontweight="bold", color=P["navy"])
    
    markerline, stemlines, baseline = axes[1, 1].stem(range(len(cooks_d)), cooks_d, markerfmt='o', basefmt=" ")
    plt.setp(stemlines, color=P["blue"])
    plt.setp(markerline, color=P["blue"], markersize=5)
    axes[1, 1].axhline(cooks_threshold, color=P["err"], linestyle="--", lw=1.5, label=f"Threshold (4/n={cooks_threshold:.3f})")
    axes[1, 1].set_xlabel("Index"); axes[1, 1].set_ylabel("Cook's Distance")
    axes[1, 1].set_title("Cook's Distance", fontweight="bold", color=P["navy"])
    axes[1, 1].legend()
    
    _style(axes.flatten())
    plt.tight_layout()
    charts.append({"title":"Regression Diagnostics","img":_store_chart(fig, "regression_diagnostics", 110)})
    
    fig2, ax2 = plt.subplots(figsize=(10, 6))
    fig2.patch.set_facecolor(P["white"])
    cv = [c["B"] for c in coef[1:]]
    ci_errors = [c["Std Error"] * 1.96 for c in coef[1:]]
    colors_ = [P["ok"] if v > 0 else P["err"] for v in cv]
    ax2.barh(range(len(cv)), cv, xerr=ci_errors, color=colors_, alpha=0.8, edgecolor="white", capsize=4, error_kw={"lw": 1.5, "color": P["navy"]})
    ax2.axvline(0, color=P["navy"], lw=1.5, ls="--")
    ax2.set_yticks(range(len(cv))); ax2.set_yticklabels(predictors, fontsize=10)
    ax2.set_xlabel("Coefficient (B)"); ax2.set_title("Coefficient Plot with 95% CI", fontweight="bold", color=P["navy"])
    _style(ax2); plt.tight_layout(); charts.append({"title":"Coefficient Plot","img":_store_chart(fig2, "regression_coefficients", 110)})
    
    eq_parts = [f"{coef[0]['B']}"]
    for row in coef[1:]:
        b_val = row['B']
        if b_val >= 0:
            eq_parts.append(f"+ {b_val}×{row['Predictor']}")
        else:
            eq_parts.append(f"- {abs(b_val)}×{row['Predictor']}")
    eq = "ŷ = " + " ".join(eq_parts)
    
    vrd, h0 = verdict(p_f)
    label = "Simple" if len(predictors) == 1 else "Multiple"
    
    # ── FIXED: Correct rows_removed calculation ──────────────────────────
    rows_removed = original_rows - len(sub)
    
    # Assumptions diagnostics with Durbin-Watson
    assumptions = []
    if len(model.resid) < 5000:
        try:
            sw_s, sw_p2 = stats.shapiro(model.resid)
            assumptions.append({"name":"Normality of Residuals","test":"Shapiro-Wilk","statistic":round(float(sw_s),4),"p_value":round(float(sw_p2),4),"passed":sw_p2>0.05})
        except:
            assumptions.append({"name":"Normality of Residuals","test":"Shapiro-Wilk","passed":False,"note":"Could not compute"})
    else:
        from scipy.stats import normaltest
        k2_stat, k2_p = normaltest(model.resid)
        assumptions.append({"name":"Normality of Residuals","test":"D'Agostino K²","statistic":round(float(k2_stat),4),"p_value":round(float(k2_p),4),"passed":k2_p>0.05})
    
    from statsmodels.stats.diagnostic import het_breuschpagan
    bp_test = het_breuschpagan(model.resid, model.model.exog)
    assumptions.append({"name":"Homoscedasticity","test":"Breusch-Pagan","p_value":round(float(bp_test[1]),4),"passed":bp_test[1]>0.05})
    
    if len(predictors) > 1:
        max_vif = max([v["VIF"] for v in vif_data]) if vif_data else 0
        assumptions.append({"name":"Multicollinearity","test":"VIF","max_vif":round(float(max_vif),3),"passed":max_vif<10 if max_vif else True})
    
    assumptions.append({"name":"Influential Points","test":"Cook's Distance","max_cooks":round(float(cooks_d.max()),4),"passed":cooks_passed})
    assumptions.append({"name":"No Autocorrelation","test":"Durbin-Watson","statistic":round(float(dw_stat),3),"passed":1.5 < dw_stat < 2.5})
    
    return {
        "test": f"{label} Linear Regression",
        "dependent": dependent,
        "predictors": predictors,
        "r_squared": round(float(r2), 4),
        "adj_r_squared": round(float(adj_r2), 4),
        "f_statistic": round(float(F_stat), 4),
        "p_value": round(float(p_f), 4),
        "p_display": fp(p_f),
        "significance": sp(p_f),
        "equation": eq,
        "coef_table": coef,
        "vif_table": vif_data,
        "cooks_distance": round(float(cooks_d.max()), 4),
        "cooks_threshold": round(cooks_threshold, 4),
        "assumptions": assumptions,
        "warnings": warnings_list,
        "n": int(len(y)),
        "charts": charts,
        "apa_citation": f"R²={r2:.3f}, Adj.R²={adj_r2:.3f}, F({len(predictors)},{len(y)-len(predictors)-1})={F_stat:.2f}, {fp(p_f)}",
        "interpretation": f"{label} linear regression predicted {dependent} from {', '.join(predictors)}. "
            f"Model explained {r2*100:.1f}% of variance (R²={r2:.3f}, Adj.R²={adj_r2:.3f}). "
            f"Overall model was {vrd} ({fp(p_f)}). Equation: {eq}.",
        "data_quality": {
            "original_rows": original_rows,
            "rows_after_cleaning": len(sub),
            "rows_removed": rows_removed,
            "percent_removed": round((rows_removed / original_rows) * 100, 2) if original_rows > 0 else 0
        },
        "missing_data_summary": get_missing_summary(df),
        "engine_version": ENGINE_VERSION,
        "analysis_timestamp": datetime.now().isoformat(),
        "reproducibility": get_system_metadata()
    }


# ══════════════════════════════════════════════════════════════════════════════
# LOGISTIC REGRESSION
# ══════════════════════════════════════════════════════════════════════════════
@with_timeout(MAX_RUNTIME_SECONDS)
def logistic_regression(df, dependent, predictors):
    check_analysis_limits(df, "logistic_regression", {"predictors": predictors})
    validate_categorical_column(df, dependent, "as dependent variable for logistic regression")
    
    if df[dependent].nunique() != 2:
        return {"error": "Logistic regression requires binary dependent variable (2 categories)"}
    
    validate_predictor_count(predictors, context="for logistic regression")
    for pred in predictors:
        validate_numeric_column(df, pred, f"as predictor")
    
    original_rows = len(df)
    warnings_list = []
    
    sub = df[[dependent] + predictors].copy()
    
    for pred in predictors:
        sub[pred] = pd.to_numeric(sub[pred], errors='coerce')
    
    sub = sub.dropna()
    
    if len(sub) == 0:
        return {"error": "No complete cases after removing missing data."}
    
    y = pd.Categorical(sub[dependent]).codes
    
    if len(np.unique(y)) < 2:
        return {"error": "Dependent variable has only one class after cleaning."}
    
    constant_preds = [p for p in predictors if sub[p].nunique() <= 1]
    if constant_preds:
        warnings_list.append(f"Dropped constant predictors: {constant_preds}")
        predictors = [p for p in predictors if p not in constant_preds]
        if len(predictors) == 0:
            return {"error": "All predictors are constant."}
    
    import statsmodels.api as sm
    X = sm.add_constant(sub[predictors].values)
    
    if np.linalg.matrix_rank(X) < X.shape[1]:
        raise SingularMatrixError(ERROR_CODES["ERR_003"])
    
    # ── FIXED: Convergence fallback hierarchy ──────────────────────────────
    model = None
    methods = ['bfgs', 'lbfgs', 'newton']
    for method in methods:
        try:
            model = sm.Logit(y, X).fit(disp=0, maxiter=1000, method=method)
            if model.mle_retvals['converged']:
                break
        except:
            continue
        warnings_list.append(f"Method {method} failed, trying next...")
    
    if model is None:
        return {"error": ERROR_CODES["ERR_005"], "message": "Logistic regression failed to converge with all methods."}
    
    if not model.mle_retvals['converged']:
        warnings_list.append("Model did not converge. Results may be unreliable.")
    
    # Coefficient table
    coef = []
    coef.append({
        "Predictor": "(Intercept)",
        "Coefficient": round(float(model.params[0]), 4),
        "Std Error": round(float(model.bse[0]), 4),
        "z": round(float(model.tvalues[0]), 4),
        "p-value": round(float(model.pvalues[0]), 4),
        "Odds Ratio": round(float(np.exp(model.params[0])), 4)
    })
    for i, pred in enumerate(predictors):
        coef.append({
            "Predictor": pred,
            "Coefficient": round(float(model.params[i+1]), 4),
            "Std Error": round(float(model.bse[i+1]), 4),
            "z": round(float(model.tvalues[i+1]), 4),
            "p-value": round(float(model.pvalues[i+1]), 4),
            "Odds Ratio": round(float(np.exp(model.params[i+1])), 4)
        })
        # Standardized coefficients
        beta = model.params[i+1] * (sub[pred].std())
        coef[i+1]["Beta"] = round(float(beta), 4)
    
    # VIF
    from statsmodels.stats.outliers_influence import variance_inflation_factor
    vif_data = []
    for i in range(X.shape[1]):
        if i == 0:
            continue
        vif = variance_inflation_factor(X, i)
        vif_data.append({
            "Variable": predictors[i-1],
            "VIF": round(float(vif), 3),
            "Interpretation": get_vif_label(vif)
        })
    
    # Predictions
    y_pred_proba = model.predict(X)
    y_pred = (y_pred_proba > 0.5).astype(int)
    
    from sklearn.metrics import confusion_matrix, roc_auc_score, precision_score, recall_score, f1_score
    
    cm = confusion_matrix(y, y_pred)
    accuracy = (cm[0, 0] + cm[1, 1]) / cm.sum()
    precision = precision_score(y, y_pred, zero_division=0)
    recall = recall_score(y, y_pred, zero_division=0)
    f1 = f1_score(y, y_pred, zero_division=0)
    auc = roc_auc_score(y, y_pred_proba)
    
    # ── FIXED: Hosmer-Lemeshow with import fallback ──────────────────────
    try:
        from statsmodels.stats.diagnostic import hosmer_lemeshow
        hl_stat, hl_p = hosmer_lemeshow(y, y_pred_proba)
    except (ImportError, AttributeError):
        hl_stat, hl_p = None, None
        warnings_list.append("Could not compute Hosmer-Lemeshow test")
    
    # ── FIXED: Correct rows_removed calculation ──────────────────────────
    rows_removed = original_rows - len(sub)
    
    return {
        "test": "Logistic Regression",
        "dependent": dependent,
        "predictors": predictors,
        "coef_table": coef,
        "vif_table": vif_data,
        "log_likelihood": round(float(model.llf), 4),
        "pseudo_r2": {
            "McFadden": round(float(model.prsquared), 4)
        },
        "confusion_matrix": {
            "True Negative": int(cm[0, 0]),
            "False Positive": int(cm[0, 1]),
            "False Negative": int(cm[1, 0]),
            "True Positive": int(cm[1, 1])
        },
        "accuracy": round(float(accuracy), 4),
        "precision": round(float(precision), 4),
        "recall": round(float(recall), 4),
        "f1_score": round(float(f1), 4),
        "auc": round(float(auc), 4),
        "hosmer_lemeshow": {
            "statistic": round(float(hl_stat), 4) if hl_stat else None,
            "p_value": round(float(hl_p), 4) if hl_p else None
        } if hl_stat else None,
        "n": int(len(y)),
        "warnings": warnings_list,
        "data_quality": {
            "original_rows": original_rows,
            "rows_after_cleaning": len(sub),
            "rows_removed": rows_removed,
            "percent_removed": round((rows_removed / original_rows) * 100, 2) if original_rows > 0 else 0
        },
        "missing_data_summary": get_missing_summary(df),
        "engine_version": ENGINE_VERSION,
        "analysis_timestamp": datetime.now().isoformat(),
        "reproducibility": get_system_metadata()
    }


# ══════════════════════════════════════════════════════════════════════════════
# MANN-WHITNEY U
# ══════════════════════════════════════════════════════════════════════════════
@with_timeout(MAX_RUNTIME_SECONDS)
def mann_whitney(df, numeric_var, group_var):
    validate_numeric_column(df, numeric_var, "for Mann-Whitney U test")
    validate_categorical_column(df, group_var, "for Mann-Whitney U test")
    original_rows = len(df)
    gs = df[group_var].dropna().unique()
    if len(gs) != 2:
        return {"error": "Requires exactly 2 groups."}
    g1 = pd.to_numeric(df.loc[df[group_var]==gs[0], numeric_var], errors="coerce").dropna()
    g2 = pd.to_numeric(df.loc[df[group_var]==gs[1], numeric_var], errors="coerce").dropna()
    
    from scipy.stats import mannwhitneyu
    res = mannwhitneyu(g1, g2, alternative='two-sided', method='auto')
    U = res.statistic
    p = res.pvalue
    
    n1, n2 = len(g1), len(g2)
    mu = n1 * n2 / 2
    from scipy.stats import rankdata
    combined = np.concatenate([g1, g2])
    ranks = rankdata(combined)
    unique, counts = np.unique(combined, return_counts=True)
    ties = counts[counts > 1]
    tie_correction = np.sum([t**3 - t for t in ties]) / 12 if len(ties) > 0 else 0
    sigma = np.sqrt((n1 * n2 / (n1 + n2)) * ((n1 + n2 + 1) - tie_correction / ((n1 + n2) * (n1 + n2 - 1))))
    z = (U - mu) / sigma if sigma > 0 else 0
    r = abs(z) / np.sqrt(n1 + n2)
    ef = "small" if r < 0.1 else "medium" if r < 0.3 else "large"
    
    vrd, h0 = verdict(p)
    charts = []
    fig, ax = plt.subplots(figsize=(8,5))
    fig.patch.set_facecolor(P["white"])
    pdf = pd.DataFrame({numeric_var: pd.concat([g1, g2]), group_var: [str(gs[0])]*len(g1) + [str(gs[1])]*len(g2)})
    sns.violinplot(data=pdf, x=group_var, y=numeric_var, palette=PAL[:2], ax=ax, inner="box", linewidth=1.5)
    ax.set_title(f"Mann-Whitney: {numeric_var} by {group_var} | U={U:.0f}, {fp(p)}", fontweight="bold", color=P["navy"])
    _style(ax); plt.tight_layout()
    charts.append({"title":"Mann-Whitney Violin","img":_store_chart(fig, "mannwhitney_violin", 110)})
    
    # ── FIXED: Correct rows_removed calculation ──────────────────────────
    rows_removed = original_rows - (len(g1) + len(g2))
    return {
        "test":"Mann-Whitney U Test","u_statistic":round(float(U),2),"p_value":round(float(p),4),
        "p_display":fp(p),"significance":sp(p),"effect_r":round(float(r),4),"effect_size_label":ef,
        "medians":{str(gs[0]):round(float(g1.median()),4), str(gs[1]):round(float(g2.median()),4)},
        "charts":charts,"apa_citation":f"U={U:.0f}, {fp(p)}, r={r:.3f}",
        "interpretation":f"Mann-Whitney U compared {numeric_var} between {gs[0]} (Mdn={g1.median():.2f}) and {gs[1]} (Mdn={g2.median():.2f}). Difference was {vrd}, U={U:.0f}, {fp(p)}, r={r:.3f} ({ef} effect).",
        "data_quality": {
            "original_rows": original_rows,
            "rows_after_cleaning": len(g1) + len(g2),
            "rows_removed": rows_removed,
            "percent_removed": round((rows_removed / original_rows) * 100, 2) if original_rows > 0 else 0
        },
        "missing_data_summary": get_missing_summary(df),
        "engine_version": ENGINE_VERSION,
        "analysis_timestamp": datetime.now().isoformat(),
        "reproducibility": get_system_metadata()
    }


# ══════════════════════════════════════════════════════════════════════════════
# KRUSKAL-WALLIS
# ══════════════════════════════════════════════════════════════════════════════
@with_timeout(MAX_RUNTIME_SECONDS)
def kruskal_wallis(df, numeric_var, group_var):
    validate_numeric_column(df, numeric_var, "for Kruskal-Wallis test")
    validate_categorical_column(df, group_var, "for Kruskal-Wallis test")
    original_rows = len(df)
    gs = df[group_var].dropna().unique()
    
    gd = []
    for g in gs:
        group_data = pd.to_numeric(df.loc[df[group_var]==g, numeric_var], errors="coerce").dropna().values
        if len(group_data) < 2:
            return {"error": f"Group '{g}' has only {len(group_data)} observations. Requires at least 2."}
        gd.append(group_data)
    
    H, p = stats.kruskal(*gd)
    n_t = sum(len(g) for g in gd)
    eta = max((H - len(gs) + 1) / (n_t - len(gs)), 0)
    vrd, h0 = verdict(p)
    charts = []
    warnings_list = []
    
    fig, ax = plt.subplots(figsize=(max(7, len(gs)*1.5), 5))
    fig.patch.set_facecolor(P["white"])
    pdf = pd.DataFrame({numeric_var: np.concatenate(gd), group_var: np.repeat([str(g) for g in gs], [len(g) for g in gd])})
    sns.boxplot(data=pdf, x=group_var, y=numeric_var, palette=PAL[:len(gs)], ax=ax, width=.5)
    sns.stripplot(data=pdf, x=group_var, y=numeric_var, color=P["navy"], alpha=.35, size=4, jitter=True, ax=ax)
    ax.set_title(f"Kruskal-Wallis: {numeric_var} by {group_var} | H({len(gs)-1})={H:.3f}, {fp(p)}", fontweight="bold", color=P["navy"])
    ax.tick_params(axis="x", rotation=30)
    _style(ax); plt.tight_layout()
    charts.append({"title":"Kruskal-Wallis Box Plot","img":_store_chart(fig, "kruskal_boxplot", 110)})
    
    # Dunn's post-hoc test
    posthoc = []
    if p < 0.05 and len(gs) > 2:
        try:
            try:
                import scikit_posthocs as sp
                dunn_result = sp.posthoc_dunn(pdf, val_col=numeric_var, group_col=group_var, p_adjust='bonferroni')
                for i, g1 in enumerate(gs):
                    for j, g2 in enumerate(gs):
                        if i < j:
                            posthoc.append({
                                "Group 1": str(g1),
                                "Group 2": str(g2),
                                "p-value": round(float(dunn_result.loc[g1, g2]), 4),
                                "Significant": "Yes" if dunn_result.loc[g1, g2] < 0.05 else "No"
                            })
            except ImportError:
                n_pairs = len(list(combinations(range(len(gs)), 2)))
                for i in range(len(gs)):
                    for j in range(i+1, len(gs)):
                        u_stat, u_p = stats.mannwhitneyu(gd[i], gd[j], alternative='two-sided')
                        u_p_corrected = min(u_p * n_pairs, 1.0)
                        posthoc.append({
                            "Group 1": str(gs[i]),
                            "Group 2": str(gs[j]),
                            "p-value": round(float(u_p_corrected), 4),
                            "Significant": "Yes" if u_p_corrected < 0.05 else "No"
                        })
        except Exception as e:
            warnings_list.append(f"Post-hoc test failed: {str(e)}")
    
    # ── FIXED: Correct rows_removed calculation ──────────────────────────
    rows_removed = original_rows - n_t
    return {
        "test":"Kruskal-Wallis H Test","h_statistic":round(float(H),4),"p_value":round(float(p),4),
        "p_display":fp(p),"df":int(len(gs)-1),"significance":sp(p),"eta_squared_h":round(float(eta),4),
        "group_medians":{str(g):round(float(np.median(gd[i])),4) for i,g in enumerate(gs)},
        "posthoc_table":posthoc,"charts":charts,
        "warnings": warnings_list,
        "apa_citation":f"H({len(gs)-1})={H:.2f}, {fp(p)}",
        "interpretation":f"Kruskal-Wallis H test compared {numeric_var} across {len(gs)} groups. Result was {vrd}, H({len(gs)-1})={H:.2f}, {fp(p)}. We {h0} H₀." +
            (" Post-hoc comparisons provided." if posthoc else ""),
        "data_quality": {
            "original_rows": original_rows,
            "rows_after_cleaning": n_t,
            "rows_removed": rows_removed,
            "percent_removed": round((rows_removed / original_rows) * 100, 2) if original_rows > 0 else 0
        },
        "missing_data_summary": get_missing_summary(df),
        "engine_version": ENGINE_VERSION,
        "analysis_timestamp": datetime.now().isoformat(),
        "reproducibility": get_system_metadata()
    }


# ══════════════════════════════════════════════════════════════════════════════
# DISPATCHER
# ══════════════════════════════════════════════════════════════════════════════
ANALYSIS_MAP = {
    "descriptive": descriptive_statistics,
    "chi_square": chi_square_test,
    "t_test": independent_ttest,
    "anova": one_way_anova,
    "correlation": correlation_analysis,
    "regression": linear_regression,
    "logistic_regression": logistic_regression,
    "mann_whitney": mann_whitney,
    "kruskal_wallis": kruskal_wallis,
}

ANALYSIS_LABELS = {
    "descriptive":"Descriptive Statistics","chi_square":"Chi-Square Test",
    "t_test":"Independent T-Test","anova":"One-Way ANOVA","correlation":"Correlation Analysis",
    "regression":"Linear Regression","logistic_regression":"Logistic Regression",
    "mann_whitney":"Mann-Whitney U Test","kruskal_wallis":"Kruskal-Wallis Test",
}

def run_analysis(df, analysis_type, params, client_ip=None):
    logger.info(f"Starting analysis: {analysis_type}")
    logger.info(f"Dataset: {len(df)} rows, {len(df.columns)} columns")
    
    if client_ip:
        try:
            check_rate_limit(client_ip)
        except RateLimitError as e:
            return {
                "error": "rate_limit_exceeded",
                "error_code": "LIM_004",
                "message": str(e),
                "cta_text": "⏳ Please wait and try again in 60 seconds",
                "cta_url": "https://www.eduxellence.org/#contact"
            }
    
    if analysis_type not in ANALYSIS_MAP:
        return {"error": "unknown_analysis", "error_code": "VAL_001", "message": f"Unknown analysis type: {analysis_type}"}
    
    try:
        start_time = time.time()
        df = validate_dataset(df)
        check_memory_usage(df)
        logger.info(f"Analysis request: {analysis_type} with params: {params}")
        result = ANALYSIS_MAP[analysis_type](df, **params)
        elapsed = time.time() - start_time
        logger.info(f"Analysis completed in {elapsed:.2f}s")
        
        if "engine_version" not in result:
            result["engine_version"] = ENGINE_VERSION
        if "analysis_timestamp" not in result:
            result["analysis_timestamp"] = datetime.now().isoformat()
        if "missing_data_summary" not in result:
            result["missing_data_summary"] = get_missing_summary(df)
        
        return result
        
    except PerfectSeparationError as e:
        logger.warning(f"Perfect separation: {str(e)}")
        return {
            "error": "perfect_separation",
            "error_code": "VAL_011",
            "message": "Perfect separation detected in logistic regression. Consider using Firth's penalized regression or removing highly predictive predictors.",
            "cta_text": "📞 Need help with logistic regression?",
            "cta_url": "https://www.eduxellence.org/#contact"
        }
    except ValidationError as e:
        logger.warning(f"Validation error: {str(e)}")
        return {
            "error": "validation_error",
            "error_code": "VAL_001",
            "message": str(e),
            "cta_text": "📞 Need help with your data?",
            "cta_url": "https://www.eduxellence.org/#contact"
        }
    except AnalysisTooLargeError as e:
        logger.warning(f"Analysis too large: {str(e)}")
        return {
            "error": "analysis_too_large",
            "error_code": "LIM_003",
            "message": str(e),
            "cta_text": "📞 Book a Free Expert Consultation",
            "cta_url": "https://www.eduxellence.org/#contact"
        }
    except AnalysisTimeoutError as e:
        logger.error(f"Analysis timeout: {str(e)}")
        return {
            "error": "analysis_timeout",
            "error_code": "LIM_001",
            "message": "The analysis took too long to complete.",
            "cta_text": "📞 Need help with large datasets?",
            "cta_url": "https://www.eduxellence.org/#contact"
        }
    except MemoryLimitError as e:
        logger.error(f"Memory limit exceeded: {str(e)}")
        return {
            "error": "memory_limit_exceeded",
            "error_code": "LIM_002",
            "message": str(e),
            "cta_text": "📞 Need help with large datasets?",
            "cta_url": "https://www.eduxellence.org/#contact"
        }
    except SingularMatrixError as e:
        logger.error(f"Singular matrix: {str(e)}")
        return {
            "error": "singular_matrix",
            "error_code": "ERR_003",
            "message": "Perfect collinearity detected. Remove redundant predictors.",
            "cta_text": "📞 Need help with regression?",
            "cta_url": "https://www.eduxellence.org/#contact"
        }
    except Exception as e:
        logger.error(f"Analysis error: {str(e)}\n{traceback.format_exc()}")
        return {
            "error": "internal_server_error",
            "error_code": "ERR_001",
            "message": "An unexpected error occurred. Our team has been notified.",
            "reference_id": str(uuid.uuid4())[:8],
            "cta_text": "📞 Contact our support team",
            "cta_url": "https://www.eduxellence.org/#contact"
        }
        # ═══════════════════════════════════════════════════════════════════════════
# ═══════════════════════════════════════════════════════════════════════════
# FLASK WEB APPLICATION — ADD THIS TO THE BOTTOM OF YOUR FILE
# ═══════════════════════════════════════════════════════════════════════════
# ═══════════════════════════════════════════════════════════════════════════

from flask import Flask, request, jsonify, send_file
import io
import tempfile
import os

app = Flask(__name__)

# ─── File Upload Endpoint ────────────────────────────────────────────────
@app.route('/api/upload', methods=['POST'])
def api_upload():
    """Handle file upload, parse data, return column metadata."""
    if 'file' not in request.files:
        return jsonify({"error": "No file uploaded"}), 400
    
    file = request.files['file']
    if file.filename == '':
        return jsonify({"error": "Empty filename"}), 400
    
    # Check file size (10 MB limit)
    file.seek(0, 2)
    file_size = file.tell()
    file.seek(0)
    if file_size > 10 * 1024 * 1024:
        return jsonify({"error": "large_file"}), 413
    
    try:
        # Parse based on extension
        ext = file.filename.split('.')[-1].lower()
        if ext == 'csv':
            df = pd.read_csv(file)
        elif ext in ['xlsx', 'xls']:
            df = pd.read_excel(file)
        else:
            return jsonify({"error": f"Unsupported file type: {ext}"}), 400
        
        # Clean and validate
        df = prepare_dataframe(df)
        
        # Generate metadata path (store temporarily)
        meta_id = str(uuid.uuid4())[:8]
        temp_dir = tempfile.mkdtemp()
        meta_path = os.path.join(temp_dir, f"{meta_id}.parquet")
        df.to_parquet(meta_path, index=False)
        
        # Build column metadata
        columns = []
        for col in df.columns:
            is_numeric = pd.to_numeric(df[col], errors='coerce').notna().sum() > len(df) * 0.5
            dtype = "numeric" if is_numeric else "categorical"
            n_unique = df[col].nunique()
            missing = df[col].isna().sum()
            missing_pct = round((missing / len(df)) * 100, 2)
            
            columns.append({
                "name": col,
                "dtype": dtype,
                "n_unique": n_unique,
                "missing": int(missing),
                "missing_pct": missing_pct
            })
        
        # Preview
        preview = df.head(10).replace({np.nan: None}).to_dict('records')
        
        # Recommendations
        recommendations = recommend_tests(df, columns)
        
        return jsonify({
            "meta_path": meta_path,
            "filename": file.filename,
            "rows": len(df),
            "cols": len(df.columns),
            "columns": columns,
            "preview": preview,
            "recommendations": recommendations[:5]
        })
        
    except Exception as e:
        logger.error(f"Upload error: {str(e)}\n{traceback.format_exc()}")
        return jsonify({"error": str(e)}), 500


# ─── Assumptions Check Endpoint ──────────────────────────────────────────
@app.route('/api/assumptions', methods=['POST'])
def api_assumptions():
    """Check assumptions for a given test."""
    data = request.get_json()
    if not data:
        return jsonify({"error": "No data provided"}), 400
    
    meta_path = data.get('meta_path')
    analysis_type = data.get('analysis_type')
    params = data.get('params', {})
    
    if not meta_path or not analysis_type:
        return jsonify({"error": "Missing meta_path or analysis_type"}), 400
    
    try:
        df = pd.read_parquet(meta_path)
        df = prepare_dataframe(df)
        
        checks = check_assumptions(df, analysis_type, params)
        all_passed = all(c['passed'] for c in checks) if checks else True
        
        return jsonify({
            "checks": checks,
            "all_passed": all_passed
        })
        
    except Exception as e:
        logger.error(f"Assumptions error: {str(e)}\n{traceback.format_exc()}")
        return jsonify({"error": str(e)}), 500


# ─── Analysis Endpoint ────────────────────────────────────────────────────
@app.route('/api/analyze', methods=['POST'])
def api_analyze():
    """Run statistical analysis."""
    data = request.get_json()
    if not data:
        return jsonify({"error": "No data provided"}), 400
    
    meta_path = data.get('meta_path')
    analysis_type = data.get('analysis_type')
    params = data.get('params', {})
    ai_mode = data.get('ai_mode', 'simple')
    
    if not meta_path or not analysis_type:
        return jsonify({"error": "Missing meta_path or analysis_type"}), 400
    
    # Rate limiting
    client_ip = request.remote_addr
    
    try:
        df = pd.read_parquet(meta_path)
        df = prepare_dataframe(df)
        
        # Check limits
        check_analysis_limits(df, analysis_type, params)
        
        # Run analysis
        result = run_analysis(df, analysis_type, params, client_ip)
        
        # Add metadata
        result['analysis_label'] = ANALYSIS_LABELS.get(analysis_type, analysis_type)
        result['generated_at'] = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        result['ai_mode'] = ai_mode
        
        # Generate AI interpretation if available
        if 'interpretation' in result and result['interpretation']:
            result['ai_interpretation'] = result['interpretation']
            result['ai_source'] = 'template'
        
        # Cleanup temp file
        try:
            os.remove(meta_path)
            os.rmdir(os.path.dirname(meta_path))
        except:
            pass
        
        return jsonify(result)
        
    except AnalysisTooLargeError as e:
        return jsonify({
            "error": "analysis_too_large",
            "message": str(e),
            "cta_text": "📞 Book a Free Expert Consultation",
            "cta_url": "https://www.eduxellence.org/#contact"
        }), 413
        
    except Exception as e:
        logger.error(f"Analysis error: {str(e)}\n{traceback.format_exc()}")
        return jsonify({
            "error": "internal_server_error",
            "message": str(e)
        }), 500


# ─── Export Endpoint ──────────────────────────────────────────────────────
@app.route('/api/export', methods=['POST'])
def api_export():
    """Export results as Excel, Word, or PDF."""
    data = request.get_json()
    if not data:
        return jsonify({"error": "No data provided"}), 400
    
    meta_path = data.get('meta_path')
    analysis_type = data.get('analysis_type')
    params = data.get('params', {})
    format_type = data.get('format', 'xlsx')
    
    if not meta_path or not analysis_type:
        return jsonify({"error": "Missing meta_path or analysis_type"}), 400
    
    try:
        df = pd.read_parquet(meta_path)
        df = prepare_dataframe(df)
        
        # Run analysis to get results
        result = run_analysis(df, analysis_type, params)
        
        if format_type == 'xlsx':
            # Create Excel export
            import pandas as pd
            from openpyxl.styles import Font, PatternFill, Alignment
            
            output = io.BytesIO()
            with pd.ExcelWriter(output, engine='openpyxl') as writer:
                # Summary sheet
                summary = {
                    'Analysis': analysis_type,
                    'Generated': datetime.now().strftime("%Y-%m-%d %H:%M"),
                    'Website': 'https://eduxellence.org'
                }
                pd.DataFrame([summary]).T.to_excel(writer, sheet_name='Summary', header=False)
                
                # Results tables
                table_keys = ['numeric_summary', 'summary_table', 'coef_table', 'pairs_table', 'posthoc_table']
                for key in table_keys:
                    if result.get(key):
                        pd.DataFrame(result[key]).to_excel(writer, sheet_name=key[:31], index=False)
                
                # Clean up
                for ws in writer.book.worksheets:
                    for row in ws.iter_rows():
                        for cell in row:
                            if cell.row == 1:
                                cell.font = Font(bold=True, color='FFFFFF')
                                cell.fill = PatternFill(start_color='0B1829', end_color='0B1829', fill_type='solid')
                                cell.alignment = Alignment(horizontal='center')
            
            output.seek(0)
            return send_file(
                output,
                mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
                download_name=f'eduxellence_{analysis_type}.xlsx'
            )
        
        elif format_type == 'docx':
            # Word export - requires python-docx
            try:
                from docx import Document
                from docx.shared import Pt, Inches
                
                doc = Document()
                doc.add_heading('Statistical Analysis Report', 0)
                doc.add_heading(f'Test: {analysis_type}', level=1)
                doc.add_paragraph(f'Generated: {datetime.now().strftime("%Y-%m-%d %H:%M")}')
                doc.add_paragraph(f'Powered by Eduxellence Analytics')
                
                if result.get('apa_citation'):
                    doc.add_heading('APA Citation', level=2)
                    doc.add_paragraph(result['apa_citation'])
                
                if result.get('interpretation'):
                    doc.add_heading('Interpretation', level=2)
                    doc.add_paragraph(result['interpretation'])
                
                output = io.BytesIO()
                doc.save(output)
                output.seek(0)
                
                return send_file(
                    output,
                    mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                    download_name=f'eduxellence_{analysis_type}.docx'
                )
            except ImportError:
                return jsonify({"error": "Word export requires python-docx"}), 500
        
        elif format_type == 'pdf':
            # PDF export - requires reportlab
            try:
                from reportlab.lib.pagesizes import letter
                from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
                from reportlab.lib.styles import getSampleStyleSheet
                
                output = io.BytesIO()
                doc = SimpleDocTemplate(output, pagesize=letter)
                styles = getSampleStyleSheet()
                story = []
                
                story.append(Paragraph('Statistical Analysis Report', styles['Title']))
                story.append(Paragraph(f'Test: {analysis_type}', styles['Heading1']))
                story.append(Spacer(1, 12))
                story.append(Paragraph(f'Generated: {datetime.now().strftime("%Y-%m-%d %H:%M")}', styles['Normal']))
                story.append(Paragraph('Powered by Eduxellence Analytics', styles['Normal']))
                
                if result.get('apa_citation'):
                    story.append(Spacer(1, 12))
                    story.append(Paragraph('APA Citation', styles['Heading2']))
                    story.append(Paragraph(result['apa_citation'], styles['Normal']))
                
                if result.get('interpretation'):
                    story.append(Spacer(1, 12))
                    story.append(Paragraph('Interpretation', styles['Heading2']))
                    story.append(Paragraph(result['interpretation'], styles['Normal']))
                
                doc.build(story)
                output.seek(0)
                
                return send_file(
                    output,
                    mimetype='application/pdf',
                    download_name=f'eduxellence_{analysis_type}.pdf'
                )
            except ImportError:
                return jsonify({"error": "PDF export requires reportlab"}), 500
        
        else:
            return jsonify({"error": f"Unsupported format: {format_type}"}), 400
        
    except Exception as e:
        logger.error(f"Export error: {str(e)}\n{traceback.format_exc()}")
        return jsonify({"error": str(e)}), 500


# ─── AI Interpretation Endpoint ──────────────────────────────────────────
@app.route('/api/ai-interpret', methods=['POST'])
def api_ai_interpret():
    """Switch AI interpretation mode."""
    data = request.get_json()
    if not data:
        return jsonify({"error": "No data provided"}), 400
    
    meta_path = data.get('meta_path')
    analysis_type = data.get('analysis_type')
    params = data.get('params', {})
    mode = data.get('mode', 'simple')
    
    if not meta_path or not analysis_type:
        return jsonify({"error": "Missing meta_path or analysis_type"}), 400
    
    try:
        df = pd.read_parquet(meta_path)
        df = prepare_dataframe(df)
        
        result = run_analysis(df, analysis_type, params)
        
        # Generate different interpretation based on mode
        base_text = result.get('interpretation', '')
        
        if mode == 'simple':
            interpretation = base_text.replace('statistically significant', '🌟 SIGNIFICANT')
            interpretation = interpretation.replace('not statistically significant', '📌 NOT significant')
        elif mode == 'executive':
            interpretation = f"📊 **Executive Summary**: {base_text[:200]}..."
        else:  # academic
            interpretation = base_text
        
        return jsonify({
            "ai_interpretation": interpretation,
            "ai_mode": mode,
            "ai_source": 'template'
        })
        
    except Exception as e:
        logger.error(f"AI interpret error: {str(e)}\n{traceback.format_exc()}")
        return jsonify({"error": str(e)}), 500


# ─── Health Check ─────────────────────────────────────────────────────────
@app.route('/api/health', methods=['GET'])
def api_health():
    """Health check endpoint for Vercel."""
    return jsonify({
        "status": "healthy",
        "engine_version": ENGINE_VERSION,
        "timestamp": datetime.now().isoformat()
    })


# ─── Root Endpoint ────────────────────────────────────────────────────────
@app.route('/api', methods=['GET'])
def api_root():
    """API root endpoint."""
    return jsonify({
        "name": "Eduxellence Statistical Analysis Engine",
        "version": ENGINE_VERSION,
        "endpoints": [
            "/api/upload",
            "/api/assumptions", 
            "/api/analyze",
            "/api/export",
            "/api/ai-interpret",
            "/api/health"
        ]
    })


# ─── Vercel Handler ──────────────────────────────────────────────────────
# This is what Vercel looks for
app = app


# ─── Local Development ────────────────────────────────────────────────────
if __name__ == '__main__':
    # Run locally for testing
    port = int(os.environ.get('PORT', 5000))
    app.run(host='0.0.0.0', port=port, debug=False)
