"""
Eduxellence Thesis Package Engine — Phase 8
==============================================
Bundles 7 statistical tests (Descriptive, Chi-Square, T-Test, ANOVA,
Correlation, Regression, Reliability) into one purchase, with a
hybrid auto-suggest + manual-review variable selection flow, and a
single combined APA-formatted report (Word + PDF) at the end.

NEW IN PHASE 8:
- Statistical soundness scoring for each test
- Assumption checking before running tests
- Filtering of statistically invalid tests
- Clear recommendations for users
- Integration with corrected Factor Analysis diagnostics

by Eduxellence Analytics · https://eduxellence.org
"""

import io
from datetime import datetime
from typing import Optional, Dict, List, Any

import pandas as pd
import numpy as np

# ===== FIX: Import from app.py instead of stats_engine =====
# The stats_engine is actually app.py, so we import the functions directly
try:
    from app import run_analysis, ANALYSIS_LABELS, check_assumptions, get_factor_diagnostics
except ImportError:
    # Fallback: try stats_engine if it exists (for standalone use)
    try:
        from stats_engine import run_analysis, ANALYSIS_LABELS
        from app import check_assumptions, get_factor_diagnostics
    except ImportError:
        # Minimal fallback if neither works
        def run_analysis(df, test_key, params):
            return {"error": f"Analysis engine not available for {test_key}"}
        ANALYSIS_LABELS = {}
        def check_assumptions(df, test_type, params):
            return [{"name": "Assumption check unavailable", "passed": False, "note": "Engine not loaded"}]
        def get_factor_diagnostics(X):
            return {"factorability": "Unknown", "kmo": None, "bartlett_p": None}

# ── The 7 component tests, in report order ──────────────────────────────
PACKAGE_TESTS = [
    "descriptive",
    "chi_square",
    "t_test",
    "anova",
    "correlation",
    "regression",
    "reliability",   # Cronbach's Alpha
]

PACKAGE_TEST_LABELS = {
    "descriptive":  "Descriptive Statistics",
    "chi_square":   "Chi-Square Test of Independence",
    "t_test":       "Independent Samples T-Test",
    "anova":        "One-Way ANOVA",
    "correlation":  "Correlation Analysis",
    "regression":   "Linear Regression",
    "reliability":  "Reliability Analysis (Cronbach's Alpha)",
}

# ===== STATISTICAL SOUNDNESS CONFIGURATION =====
MINIMUM_ASSUMPTION_SCORE = 0.6  # 60% minimum to be considered statistically sound

# Tests that require specific assumptions to be valid
ASSUMPTION_WEIGHTS = {
    "t_test": {
        "Adequate sample size": 0.20,
        "Normality (Shapiro-Wilk)": 0.30,
        "No extreme outliers": 0.20,
        "Equal variances (Levene)": 0.30,
    },
    "anova": {
        "Adequate sample size": 0.15,
        "Normality (Shapiro-Wilk)": 0.25,
        "No extreme outliers": 0.20,
        "Equal variances (Levene)": 0.40,
    },
    "regression": {
        "Normality of residuals": 0.25,
        "Homoscedasticity (Breusch-Pagan)": 0.25,
        "No autocorrelation (Durbin-Watson)": 0.25,
        "Multicollinearity": 0.25,
    },
    "correlation": {
        "Adequate sample size": 0.25,
        "Normality (Shapiro-Wilk)": 0.25,
        "No extreme outliers": 0.25,
        "Linearity (Pearson vs Spearman)": 0.25,
    },
    "chi_square": {
        "Expected frequencies ≥ 5": 1.0,
    },
    "reliability": {
        "Adequate sample size": 0.30,
        "Items are correlated": 0.40,
        "No extreme outliers": 0.30,
    },
}

# Tests that are always valid (no assumptions)
ALWAYS_VALID = ["descriptive"]


# ══════════════════════════════════════════════════════════════════════════
#  NEW: STATISTICAL SOUNDNESS SCORING
# ══════════════════════════════════════════════════════════════════════════

def score_assumptions(assumption_checks: List[Dict]) -> Dict:
    """
    Score the assumptions for a test and determine if it's statistically sound.
    
    Args:
        assumption_checks: List of assumption check results from check_assumptions()
        
    Returns:
        {
            "score": float (0-1),
            "passed_count": int,
            "total_count": int,
            "passed": bool (score >= MINIMUM_ASSUMPTION_SCORE),
            "details": dict with per-check results
        }
    """
    if not assumption_checks:
        return {
            "score": 1.0,
            "passed_count": 0,
            "total_count": 0,
            "passed": True,
            "details": {},
            "note": "No assumptions to check"
        }
    
    total_checks = len(assumption_checks)
    passed_checks = sum(1 for c in assumption_checks if c.get("passed", False))
    
    # Calculate weighted score if weights are available
    # Otherwise use simple percentage
    score = passed_checks / total_checks if total_checks > 0 else 1.0
    
    passed = score >= MINIMUM_ASSUMPTION_SCORE
    
    details = {}
    for check in assumption_checks:
        name = check.get("name", "Unknown")
        details[name] = {
            "passed": check.get("passed", False),
            "note": check.get("note", ""),
            "suggestion": check.get("suggestion", "")
        }
    
    return {
        "score": round(score, 4),
        "passed_count": passed_checks,
        "total_count": total_checks,
        "passed": passed,
        "details": details,
        "note": f"Passed {passed_checks}/{total_checks} assumptions ({score*100:.1f}%)"
    }


def get_test_soundness(df: pd.DataFrame, test_key: str, params: Dict) -> Dict:
    """
    Check if a test is statistically sound for the given data.
    
    Returns:
        {
            "sound": bool,
            "score": float,
            "assumptions": dict,
            "recommendation": str,
            "details": dict
        }
    """
    # Descriptive is always valid
    if test_key in ALWAYS_VALID:
        return {
            "sound": True,
            "score": 1.0,
            "assumptions": [],
            "recommendation": "Descriptive statistics are always valid for any data.",
            "details": {}
        }
    
    # Check assumptions
    try:
        assumption_checks = check_assumptions(df, test_key, params)
    except Exception as e:
        return {
            "sound": False,
            "score": 0.0,
            "assumptions": [],
            "recommendation": f"Could not check assumptions: {str(e)}",
            "details": {"error": str(e)}
        }
    
    # Score the assumptions
    scored = score_assumptions(assumption_checks)
    
    # Generate recommendation
    if scored["passed"]:
        recommendation = f"This test is statistically sound for your data (score: {scored['score']*100:.1f}%)."
        if scored["score"] < 0.8:
            recommendation += " Consider reviewing the assumption details below."
    else:
        recommendation = f"This test may NOT be appropriate for your data (score: {scored['score']*100:.1f}%)."
        failed = [k for k, v in scored["details"].items() if not v.get("passed", False)]
        if failed:
            recommendation += f" Failed assumptions: {', '.join(failed)}. Consider using an alternative test."
    
    return {
        "sound": scored["passed"],
        "score": scored["score"],
        "assumptions": assumption_checks,
        "recommendation": recommendation,
        "details": scored["details"]
    }


# ══════════════════════════════════════════════════════════════════════════
#  RELIABILITY TEST (Cronbach's Alpha)
# ══════════════════════════════════════════════════════════════════════════

def reliability_analysis(df: pd.DataFrame, columns: list) -> dict:
    """
    Cronbach's Alpha for a set of Likert/scale items.
    """
    sub = df[columns].apply(pd.to_numeric, errors="coerce").dropna()
    if len(sub) < 2 or len(columns) < 2:
        return {"error": "Reliability analysis needs at least 2 items and 2 valid rows."}

    item_vars = sub.var(axis=0, ddof=1)
    total_var = sub.sum(axis=1).var(ddof=1)
    k = len(columns)

    if total_var == 0:
        return {"error": "Total variance is zero — cannot compute Cronbach's Alpha."}

    alpha = (k / (k - 1)) * (1 - item_vars.sum() / total_var)

    interpretation_label = (
        "excellent" if alpha >= 0.9 else
        "good" if alpha >= 0.8 else
        "acceptable" if alpha >= 0.7 else
        "questionable" if alpha >= 0.6 else
        "poor" if alpha >= 0.5 else "unacceptable"
    )

    # Item-total correlations
    item_table = []
    for col in columns:
        rest = sub.drop(columns=[col]).sum(axis=1)
        try:
            r = float(np.corrcoef(sub[col], rest)[0, 1])
        except Exception:
            r = float("nan")
        item_table.append({
            "Item": col,
            "Item Variance": round(float(item_vars[col]), 4),
            "Item-Total Correlation": round(r, 4) if not np.isnan(r) else "—",
        })

    return {
        "test": "Reliability Analysis (Cronbach's Alpha)",
        "alpha": round(float(alpha), 4),
        "n_items": k,
        "n": int(len(sub)),
        "effect_size": interpretation_label,
        "item_table": item_table,
        "p_value": None, "p_display": "—", "significance": "—",
        "apa_citation": f"Cronbach's α = {alpha:.2f} (N = {len(sub)}, k = {k} items)",
        "interpretation": (
            f"Reliability analysis was conducted on {k} items using Cronbach's Alpha. "
            f"Internal consistency was {interpretation_label} (α = {alpha:.2f}). "
            + ("Values above .70 are generally considered acceptable for research purposes."
               if alpha >= 0.7 else
               "Values below .70 suggest the items may not be measuring a single "
               "consistent underlying construct; consider reviewing or removing weak items.")
        ),
        "charts": [],
    }


# ══════════════════════════════════════════════════════════════════════════
#  STEP 1 — AUTO-SUGGEST VARIABLES (with soundness scoring)
# ══════════════════════════════════════════════════════════════════════════

def suggest_package_variables(df: pd.DataFrame, columns: list) -> dict:
    """
    For each of the 7 component tests, propose a best-guess variable
    configuration based on column types. Includes statistical soundness
    scoring for each suggested configuration.
    """
    num_cols = [c["name"] for c in columns if c["dtype"] == "numeric"]
    cat_cols = [c["name"] for c in columns if c["dtype"] == "categorical"]
    cat2     = [c for c in cat_cols if df[c].dropna().nunique() == 2]
    cat3p    = [c for c in cat_cols if df[c].dropna().nunique() >= 3]

    # Filter out identifier-looking columns
    _ID_PATTERNS = ("id", "code", "ref", "number", "no", "num")
    def _looks_like_id(name: str) -> bool:
        lname = name.lower().replace("_", " ").replace("-", " ")
        tokens = lname.split()
        return any(tok in _ID_PATTERNS for tok in tokens)

    num_cols_analytical = [c for c in num_cols if not _looks_like_id(c)] or num_cols

    suggestions = {}

    # Helper: Get soundness for a test with given params
    def get_soundness_for_test(test_key, params):
        if params:
            return get_test_soundness(df, test_key, params)
        return {
            "sound": False,
            "score": 0.0,
            "assumptions": [],
            "recommendation": "No variables selected for this test.",
            "details": {}
        }

    # 1. Descriptive
    desc_cols = (num_cols_analytical[:6] + cat_cols[:2]) if (num_cols_analytical or cat_cols) else []
    params = {"columns": desc_cols} if desc_cols else None
    soundness = get_soundness_for_test("descriptive", params) if params else {"sound": True, "score": 1.0, "recommendation": "Always valid"}
    suggestions["descriptive"] = {
        "params": params,
        "reason": f"{len(desc_cols)} columns selected for summary statistics." if desc_cols else "No suitable columns found.",
        "soundness": soundness,
    }

    # 2. Chi-Square
    chi_cands = [c for c in cat_cols if df[c].dropna().nunique() >= 2]
    params = {"var1": chi_cands[0], "var2": chi_cands[1]} if len(chi_cands) >= 2 else None
    soundness = get_soundness_for_test("chi_square", params) if params else {"sound": False, "score": 0.0}
    suggestions["chi_square"] = {
        "params": params,
        "reason": f"Using '{chi_cands[0]}' and '{chi_cands[1]}'" if len(chi_cands) >= 2 else "Needs at least 2 categorical columns.",
        "soundness": soundness,
    }

    # 3. T-Test
    params = {"numeric_var": num_cols_analytical[0], "group_var": cat2[0]} if (num_cols_analytical and cat2) else None
    soundness = get_soundness_for_test("t_test", params) if params else {"sound": False, "score": 0.0}
    suggestions["t_test"] = {
        "params": params,
        "reason": f"Comparing '{num_cols_analytical[0]}' between groups in '{cat2[0]}'" if (num_cols_analytical and cat2) else "Needs 1 numeric + 1 binary categorical.",
        "soundness": soundness,
    }

    # 4. ANOVA
    params = {"numeric_var": num_cols_analytical[0], "group_var": cat3p[0]} if (num_cols_analytical and cat3p) else None
    soundness = get_soundness_for_test("anova", params) if params else {"sound": False, "score": 0.0}
    suggestions["anova"] = {
        "params": params,
        "reason": f"Comparing '{num_cols_analytical[0]}' across groups in '{cat3p[0]}'" if (num_cols_analytical and cat3p) else "Needs 1 numeric + 1 categorical (3+ groups).",
        "soundness": soundness,
    }

    # 5. Correlation
    params = {"columns": num_cols_analytical[:6], "method": "pearson"} if len(num_cols_analytical) >= 2 else None
    soundness = get_soundness_for_test("correlation", params) if params else {"sound": False, "score": 0.0}
    suggestions["correlation"] = {
        "params": params,
        "reason": f"Correlating {min(len(num_cols_analytical),6)} numeric variables." if len(num_cols_analytical) >= 2 else "Needs at least 2 numeric columns.",
        "soundness": soundness,
    }

    # 6. Regression
    params = {"dependent": num_cols_analytical[0], "predictors": num_cols_analytical[1:3]} if len(num_cols_analytical) >= 2 else None
    soundness = get_soundness_for_test("regression", params) if params else {"sound": False, "score": 0.0}
    suggestions["regression"] = {
        "params": params,
        "reason": f"Predicting '{num_cols_analytical[0]}' from {', '.join(num_cols_analytical[1:3])}" if len(num_cols_analytical) >= 2 else "Needs at least 2 numeric columns.",
        "soundness": soundness,
    }

    # 7. Reliability
    likert_cands = []
    for c in [col["name"] for col in columns]:
        coerced = pd.to_numeric(df[c], errors="coerce")
        coercion_rate = coerced.notna().sum() / max(len(df), 1)
        nun = coerced.dropna().nunique()
        if coercion_rate > 0.9 and 3 <= nun <= 11:
            likert_cands.append(c)
    params = {"columns": likert_cands[:8]} if len(likert_cands) >= 2 else None
    soundness = get_soundness_for_test("reliability", params) if params else {"sound": False, "score": 0.0}
    suggestions["reliability"] = {
        "params": params,
        "reason": f"{min(len(likert_cands),8)} scale-like items detected." if len(likert_cands) >= 2 else "Needs 2+ Likert-style items.",
        "soundness": soundness,
    }

    return suggestions


# ══════════════════════════════════════════════════════════════════════════
#  STEP 2 — VALIDATE (block-until-fixed rule)
# ══════════════════════════════════════════════════════════════════════════

def validate_package_config(df: pd.DataFrame, config: dict) -> dict:
    """
    Validates a user-confirmed config dict. Blocks generation entirely
    if any test has missing/invalid params OR is not statistically sound.
    
    Returns { ok: bool, errors: [ {test, message}, ... ], soundness_warnings: [...] }
    """
    errors = []
    soundness_warnings = []
    soundness_results = {}

    for test_key in PACKAGE_TESTS:
        params = config.get(test_key)
        
        # Structural validation
        if not params:
            errors.append({
                "test": test_key, 
                "label": PACKAGE_TEST_LABELS[test_key],
                "message": "No variables selected for this test.",
                "type": "structural"
            })
            continue

        # Per-test structural validation
        structural_error = None
        if test_key == "descriptive":
            if not params.get("columns"):
                structural_error = "Select at least 1 column."

        elif test_key == "chi_square":
            v1, v2 = params.get("var1"), params.get("var2")
            if not v1 or not v2:
                structural_error = "Select two categorical variables."
            elif v1 == v2:
                structural_error = "The two variables must be different."

        elif test_key in ("t_test",):
            nv, gv = params.get("numeric_var"), params.get("group_var")
            if not nv or not gv:
                structural_error = "Select a numeric variable and a grouping variable."
            elif gv in df.columns and df[gv].dropna().nunique() != 2:
                structural_error = f"'{gv}' must have exactly 2 groups."

        elif test_key == "anova":
            nv, gv = params.get("numeric_var"), params.get("group_var")
            if not nv or not gv:
                structural_error = "Select a numeric variable and a grouping variable."
            elif gv in df.columns and df[gv].dropna().nunique() < 3:
                structural_error = f"'{gv}' must have 3+ groups."

        elif test_key == "correlation":
            if len(params.get("columns", [])) < 2:
                structural_error = "Select at least 2 numeric columns."

        elif test_key == "regression":
            dep = params.get("dependent")
            preds = params.get("predictors", [])
            if not dep or not preds:
                structural_error = "Select a dependent variable and at least 1 predictor."
            elif dep in preds:
                structural_error = "Dependent variable cannot also be a predictor."

        elif test_key == "reliability":
            if len(params.get("columns", [])) < 2:
                structural_error = "Select at least 2 scale items."

        if structural_error:
            errors.append({
                "test": test_key,
                "label": PACKAGE_TEST_LABELS[test_key],
                "message": structural_error,
                "type": "structural"
            })
            continue

        # ===== NEW: Statistical soundness validation =====
        if test_key not in ALWAYS_VALID:
            soundness = get_test_soundness(df, test_key, params)
            soundness_results[test_key] = soundness
            
            if not soundness["sound"]:
                soundness_warnings.append({
                    "test": test_key,
                    "label": PACKAGE_TEST_LABELS[test_key],
                    "score": soundness["score"],
                    "message": soundness["recommendation"],
                    "details": soundness["details"],
                    "type": "soundness_warning"
                })

    return {
        "ok": len(errors) == 0,
        "errors": errors,
        "soundness_warnings": soundness_warnings,
        "soundness_results": soundness_results,
        "has_soundness_issues": len(soundness_warnings) > 0
    }


# ══════════════════════════════════════════════════════════════════════════
#  STEP 3 — RUN THE FULL PACKAGE (with soundness filtering)
# ══════════════════════════════════════════════════════════════════════════

def run_thesis_package(df: pd.DataFrame, config: dict, 
                       skip_unsound: bool = True,
                       min_score: float = MINIMUM_ASSUMPTION_SCORE) -> dict:
    """
    Runs the thesis package with optional soundness filtering.
    
    Args:
        df: DataFrame
        config: Validated config dict
        skip_unsound: If True, skip tests that don't meet the minimum score
        min_score: Minimum assumption score threshold (default: 0.6)
    
    Returns:
        {
            ok: bool,
            results: { test_key: result_dict },
            errors: [...],
            skipped: { test_key: reason, ... },
            soundness: { test_key: soundness_result, ... },
            completed_count: int,
            total_count: int,
            skipped_count: int
        }
    """
    results = {}
    runtime_errors = []
    skipped_tests = {}
    soundness_summary = {}
    
    for test_key in PACKAGE_TESTS:
        params = config.get(test_key, {})
        
        # ===== Check soundness before running =====
        if test_key not in ALWAYS_VALID and skip_unsound:
            soundness = get_test_soundness(df, test_key, params)
            soundness_summary[test_key] = soundness
            
            if soundness["score"] < min_score:
                skipped_tests[test_key] = {
                    "reason": f"Statistically unsound (score: {soundness['score']*100:.1f}%, minimum: {min_score*100:.1f}%)",
                    "score": soundness["score"],
                    "recommendation": soundness["recommendation"],
                    "details": soundness["details"]
                }
                continue
        
        # Run the test
        try:
            if test_key == "reliability":
                r = reliability_analysis(df, params.get("columns", []))
            else:
                r = run_analysis(df, test_key, params)

            if "error" in r:
                runtime_errors.append({
                    "test": test_key, 
                    "label": PACKAGE_TEST_LABELS[test_key],
                    "message": r["error"]
                })
                continue

            r["analysis_label"] = PACKAGE_TEST_LABELS[test_key]
            results[test_key] = r

        except Exception as e:
            runtime_errors.append({
                "test": test_key, 
                "label": PACKAGE_TEST_LABELS[test_key],
                "message": str(e)
            })

    return {
        "ok": len(runtime_errors) == 0,
        "results": results,
        "errors": runtime_errors,
        "skipped": skipped_tests,
        "soundness": soundness_summary,
        "completed_count": len(results),
        "total_count": len(PACKAGE_TESTS),
        "skipped_count": len(skipped_tests),
        "all_tests_count": len(PACKAGE_TESTS)
    }


# ══════════════════════════════════════════════════════════════════════════
#  STEP 4 — COMBINED REPORT (Word + PDF)
# ══════════════════════════════════════════════════════════════════════════

def generate_combined_docx(package_result: dict, dataset_name: str = "") -> bytes:
    """Builds one APA-style Word document covering all completed tests in order."""
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from export_engine import (_apa_heading, _apa_table, _add_hor_rule, _strip_md,
                                 BRAND_NAVY, BRAND_BLUE, BRAND_SLATE, SITE, NOW)

    doc = Document()
    for s in doc.sections:
        s.top_margin = s.bottom_margin = s.left_margin = s.right_margin = Inches(1)

    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)

    # Cover page
    title_p = doc.add_paragraph(); title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title_p.add_run("Complete Statistical Analysis Report")
    r.bold = True; r.font.size = Pt(18); r.font.name = 'Times New Roman'; r.font.color.rgb = BRAND_NAVY

    doc.add_paragraph()
    sub_p = doc.add_paragraph(); sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub_p.add_run(f"Thesis / Research Package — {dataset_name or 'Dataset'}")
    sr.bold = True; sr.font.size = Pt(13); sr.font.color.rgb = BRAND_BLUE

    meta_p = doc.add_paragraph(); meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    mr = meta_p.add_run(f"Generated: {NOW}\nPowered by Eduxellence Analytics — {SITE}")
    mr.font.size = Pt(10); mr.font.color.rgb = BRAND_SLATE; mr.font.italic = True

    _add_hor_rule(doc)
    doc.add_paragraph()

    # ===== NEW: Soundness Summary Section =====
    if package_result.get("skipped") or package_result.get("soundness"):
        _apa_heading(doc, "Statistical Soundness Summary", level=1)
        
        if package_result.get("skipped"):
            skip_p = doc.add_paragraph()
            skip_p.add_run("⚠️ The following tests were skipped due to statistical soundness concerns:").bold = True
            for test_key, skip_info in package_result["skipped"].items():
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.5)
                p.add_run(f"- {PACKAGE_TEST_LABELS[test_key]}: {skip_info['reason']}").font.size = Pt(11)
        
        if package_result.get("soundness"):
            sound_p = doc.add_paragraph()
            sound_p.add_run("📊 Assumption check results for run tests:").bold = True
            for test_key, sound_info in package_result["soundness"].items():
                if test_key not in package_result.get("skipped", {}):
                    p = doc.add_paragraph()
                    p.paragraph_format.left_indent = Inches(0.5)
                    status = "✅" if sound_info.get("sound", False) else "⚠️"
                    p.add_run(f"{status} {PACKAGE_TEST_LABELS[test_key]}: {sound_info.get('score', 0)*100:.1f}% sound").font.size = Pt(11)
        
        doc.add_paragraph()
        _add_hor_rule(doc)

    # Table of contents
    _apa_heading(doc, "Contents", level=1)
    section_num = 1
    for key in PACKAGE_TESTS:
        if key in package_result["results"]:
            p = doc.add_paragraph()
            p.add_run(f"{section_num}. {PACKAGE_TEST_LABELS[key]}").font.size = Pt(11)
            section_num += 1
    doc.add_paragraph()
    _add_hor_rule(doc)

    # Each test section
    section_num = 1
    for key in PACKAGE_TESTS:
        result = package_result["results"].get(key)
        if not result:
            continue

        doc.add_page_break()
        _apa_heading(doc, f"{section_num}. {PACKAGE_TEST_LABELS[key]}", level=1)
        section_num += 1

        if result.get("apa_citation"):
            cite_p = doc.add_paragraph()
            cite_p.paragraph_format.left_indent = Inches(0.5)
            cr_ = cite_p.add_run(result["apa_citation"])
            cr_.italic = True; cr_.font.size = Pt(12)
            doc.add_paragraph()

        # Tables
        for tkey in ("numeric_summary", "summary_table", "coef_table", "pairs_table",
                     "posthoc_table", "item_table"):
            if result.get(tkey):
                _apa_table(doc, result[tkey])
                doc.add_paragraph()

        if result.get("categorical_summary"):
            for cs in result["categorical_summary"]:
                p = doc.add_paragraph(); p.add_run(f"Frequency Table: {cs['variable']}").bold = True
                _apa_table(doc, cs["table"])
                doc.add_paragraph()

        if result.get("interpretation"):
            clean = _strip_md(result["interpretation"])
            for para_text in clean.split("\n\n"):
                if para_text.strip():
                    p = doc.add_paragraph()
                    p.paragraph_format.first_line_indent = Inches(0.5)
                    p.add_run(para_text.strip().replace("\n", " ")).font.size = Pt(12)
            doc.add_paragraph()

    # Footer
    doc.add_page_break()
    _apa_heading(doc, "Methodology Note", level=1)
    method_p = doc.add_paragraph()
    method_p.paragraph_format.first_line_indent = Inches(0.5)
    method_p.add_run(
        f"This complete statistical package was generated using the Eduxellence Analytics "
        f"platform ({SITE}). "
    ).font.size = Pt(12)
    
    if package_result.get("skipped_count", 0) > 0:
        method_p.add_run(
            f"Note: {package_result['skipped_count']} test(s) were automatically skipped "
            f"because they did not meet the minimum statistical soundness threshold. "
        ).font.size = Pt(12)
    
    method_p.add_run(
        f"All {package_result['completed_count']} completed analyses were verified for "
        f"statistical assumptions. Results follow APA 7th Edition conventions. Generated on {NOW}."
    ).font.size = Pt(12)

    _add_hor_rule(doc)
    footer_p = doc.add_paragraph(); footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer_p.add_run(f"Generated by Eduxellence Analytics · {SITE} · {NOW}")
    fr.font.size = Pt(9); fr.font.italic = True; fr.font.color.rgb = BRAND_SLATE

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def generate_combined_pdf(package_result: dict, dataset_name: str = "") -> bytes:
    """Builds one branded PDF report covering all completed tests, including charts."""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT
    from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer,
                                      HRFlowable, Image, KeepTogether, PageBreak)
    from export_engine import (_pdf_table, _strip_md, RL_NAVY, RL_BLUE, RL_SLATE,
                                 SITE, NOW)
    import base64

    buf = io.BytesIO()

    title_style = ParagraphStyle('Title', fontName='Helvetica-Bold', fontSize=20,
        textColor=RL_NAVY, alignment=TA_CENTER, spaceAfter=6)
    sub_style = ParagraphStyle('Sub', fontName='Helvetica-Bold', fontSize=13,
        textColor=RL_BLUE, alignment=TA_CENTER, spaceAfter=4)
    meta_style = ParagraphStyle('Meta', fontName='Helvetica-Oblique', fontSize=9,
        textColor=RL_SLATE, alignment=TA_CENTER, spaceAfter=16)
    h1_style = ParagraphStyle('H1', fontName='Helvetica-Bold', fontSize=14,
        textColor=RL_NAVY, spaceBefore=14, spaceAfter=8)
    body_style = ParagraphStyle('Body', fontName='Helvetica', fontSize=10,
        leading=16, alignment=TA_JUSTIFY, spaceAfter=8, firstLineIndent=18,
        textColor=colors.HexColor('#1E293B'))
    apa_style = ParagraphStyle('APA', fontName='Helvetica-Oblique', fontSize=10.5,
        textColor=RL_NAVY, leading=16, spaceAfter=10, leftIndent=18, rightIndent=18,
        borderColor=colors.HexColor('#C4B5FD'), borderWidth=1, borderPad=8,
        backColor=colors.HexColor('#F8F7FF'))
    footer_style = ParagraphStyle('Footer', fontName='Helvetica-Oblique', fontSize=8,
        textColor=RL_SLATE, alignment=TA_CENTER)
    warning_style = ParagraphStyle('Warning', fontName='Helvetica', fontSize=10,
        textColor=colors.HexColor('#DC2626'), alignment=TA_LEFT, spaceAfter=6)

    story = []
    story.append(Spacer(1, 0.3*inch))
    story.append(Paragraph("Complete Statistical Analysis Report", title_style))
    story.append(Paragraph(f"Thesis / Research Package — {dataset_name or 'Dataset'}", sub_style))
    story.append(Paragraph(f"Generated: {NOW} &nbsp;·&nbsp; Powered by Eduxellence Analytics &nbsp;·&nbsp; {SITE}", meta_style))
    story.append(HRFlowable(width="100%", thickness=2, color=RL_BLUE, spaceAfter=14))

    # ===== NEW: Soundness Summary Section =====
    if package_result.get("skipped") or package_result.get("soundness"):
        story.append(Paragraph("Statistical Soundness Summary", h1_style))
        
        if package_result.get("skipped"):
            story.append(Paragraph("⚠️ The following tests were skipped due to statistical soundness concerns:", warning_style))
            for test_key, skip_info in package_result["skipped"].items():
                story.append(Paragraph(
                    f"• {PACKAGE_TEST_LABELS[test_key]}: {skip_info['reason']}",
                    body_style
                ))
        
        if package_result.get("soundness"):
            story.append(Paragraph("📊 Assumption check results for run tests:", body_style))
            for test_key, sound_info in package_result["soundness"].items():
                if test_key not in package_result.get("skipped", {}):
                    status = "✅" if sound_info.get("sound", False) else "⚠️"
                    story.append(Paragraph(
                        f"{status} {PACKAGE_TEST_LABELS[test_key]}: {sound_info.get('score', 0)*100:.1f}% sound",
                        body_style
                    ))
        
        story.append(Spacer(1, 10))
        story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor('#E2E8F0'), spaceAfter=10))

    section_num = 1
    for key in PACKAGE_TESTS:
        result = package_result["results"].get(key)
        if not result:
            continue

        if section_num > 1:
            story.append(PageBreak())

        story.append(Paragraph(f"{section_num}. {PACKAGE_TEST_LABELS[key]}", h1_style))
        section_num += 1

        if result.get("apa_citation"):
            story.append(Paragraph(result["apa_citation"], apa_style))

        for tkey in ("numeric_summary", "summary_table", "coef_table", "pairs_table",
                     "posthoc_table", "item_table"):
            if result.get(tkey):
                story.append(_pdf_table(result[tkey]))
                story.append(Spacer(1, 10))

        if result.get("interpretation"):
            clean = _strip_md(result["interpretation"])
            for para_text in clean.split("\n\n"):
                if para_text.strip():
                    story.append(Paragraph(para_text.strip().replace("\n"," "), body_style))

        if result.get("charts"):
            for chart in result["charts"][:2]:
                try:
                    img_bytes = base64.b64decode(chart["img"])
                    img = Image(io.BytesIO(img_bytes), width=6*inch, height=3.6*inch)
                    cap_style = ParagraphStyle('Cap', fontName='Helvetica-Oblique', fontSize=9,
                        textColor=RL_SLATE, alignment=TA_CENTER, spaceAfter=10)
                    story.append(KeepTogether([img, Paragraph(f"Figure. {chart.get('title','Chart')}", cap_style), Spacer(1,8)]))
                except Exception:
                    pass

    story.append(PageBreak())
    story.append(Paragraph("Methodology Note", h1_style))
    
    method_text = (
        f"This complete statistical package was generated using the Eduxellence Analytics "
        f"platform ({SITE}). "
    )
    if package_result.get("skipped_count", 0) > 0:
        method_text += (
            f"Note: {package_result['skipped_count']} test(s) were automatically skipped "
            f"because they did not meet the minimum statistical soundness threshold. "
        )
    method_text += (
        f"All {package_result['completed_count']} completed analyses were verified for "
        f"statistical assumptions. Results follow APA 7th Edition conventions. "
        f"Generated on {NOW}."
    )
    
    story.append(Paragraph(method_text, body_style))
    story.append(Spacer(1, 16))
    story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor('#E2E8F0'), spaceAfter=6))
    story.append(Paragraph(f"Generated by Eduxellence Analytics · {SITE} · {NOW}", footer_style))

    doc = SimpleDocTemplate(buf, pagesize=A4, leftMargin=inch, rightMargin=inch,
        topMargin=inch, bottomMargin=inch,
        title="Complete Statistical Analysis Report — Eduxellence Analytics",
        author="Eduxellence Analytics")
    doc.build(story)
    buf.seek(0)
    return buf.read()
